# ======================================================================================
# ARCHIVO: Pagina_Covinoc.py (v17 - COMPLETO: Reporte Revisión TOTAL COVINOC)
# MODIFICADO:
#           1. TAB 1: El Excel "Listado_Clientes_Para_Revision.xlsx" ahora descarga
#              TODOS los clientes que existen en ReporteTransacciones (Covinoc), 
#              independientemente de si están al día, vencidos o filtrados.
#           2. Se mantiene toda la estética y funcionalidad previa.
#
# REQUISITOS (requirements.txt):
#           - streamlit
#           - pandas
#           - openpyxl
#           - xlsxwriter
#           - dropbox
#           - plotly
#           - python-docx
# ======================================================================================
import streamlit as st
import streamlit.components.v1 as components
import pandas as pd
import toml
import os
from io import BytesIO, StringIO
import plotly.express as px
import plotly.graph_objects as go
import unicodedata
import re
from datetime import datetime, timedelta
import dropbox
import glob
import urllib.parse
import urllib.request as urllib_request
import json
import zipfile

# --- LIBRERÍA PARA WORD ---
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    st.error("⚠️ Librería 'python-docx' no detectada. Por favor agrégala a requirements.txt.")

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(
    page_title="Gestión Covinoc",
    page_icon="🛡️",
    layout="wide"
)

# --- PALETA DE COLORES INSTITUCIONAL ---
PALETA_COLORES = {
    "primario": "#B21917",       # Rojo Oscuro Institucional
    "secundario": "#E73537",     # Rojo Claro
    "acento": "#F0833A",         # Naranja
    "destacado": "#F9B016",      # Amarillo
    "fondo_claro": "#FAFAFA",    # Fondo Web
    "fondo_suave": "#FEF4C0",    # Amarillo Pálido
    "texto_claro": "#FFFFFF",
    "texto_oscuro": "#31333F",
    "exito_verde": "#388E3C"
}

# Diccionario de Vendedores y Teléfonos
VENDEDORES_WHATSAPP = {
    "HUGO NELSON ZAPATA RAYO": "+573117658075",
    "TANIA RESTREPO BENJUMEA": "+573207425966",
    "DIEGO MAURICIO GARCIA RENGIFO": "+573205046277",
    "PABLO CESAR MAFLA BAÑOL": "+573103738523",
    "PEREZ SANTA GUSTAVO ADOLFO": "+573103663945",
    "ELISABETH CAROLINA IBARRA MANSO": "+573156224689",
    "CARLOS ALBERTO CASTRILLON LOPEZ": "+573147577658",
    "LEIVYN GRABIEL GARCIA MUNOZ": "+573127574279",
    "LEDUYN MELGAREJO ARIAS": "+573006620143",
    "JERSON ATEHORTUA OLARTE": "+573104952606"
}

st.markdown(f"""
<style>
    .stApp {{ background-color: {PALETA_COLORES['fondo_claro']}; }}
    .stMetric {{ 
        background-color: #FFFFFF; 
        border-radius: 10px; 
        padding: 20px; 
        border-left: 5px solid {PALETA_COLORES['primario']};
        box-shadow: 0 4px 8px rgba(0,0,0,0.05);
    }}
    .stTabs [data-baseweb="tab-list"] {{ gap: 24px; }}
    .stTabs [data-baseweb="tab"] {{ height: 50px; white-space: pre-wrap; background-color: transparent; border-radius: 4px 4px 0px 0px; border-bottom: 2px solid #C0C0C0; }}
    .stTabs [aria-selected="true"] {{ 
        border-bottom: 3px solid {PALETA_COLORES['primario']}; 
        color: {PALETA_COLORES['primario']}; 
        font-weight: bold; 
        background-color: {PALETA_COLORES['fondo_suave']};
    }}
    div[data-baseweb="input"], div[data-baseweb="select"], div[data-baseweb="text-area"] {{ 
        background-color: #FFFFFF; 
        border: 1.5px solid {PALETA_COLORES['acento']}; 
        border-radius: 8px; 
        box-shadow: 0 2px 4px rgba(0,0,0,0.1); 
        padding-left: 5px; 
    }}
    /* Botones personalizados */
    div.stButton > button:first-child {{
        background-color: {PALETA_COLORES['primario']};
        color: white;
        border: none;
        border-radius: 6px;
    }}
    div.stButton > button:hover {{
        background-color: {PALETA_COLORES['secundario']};
        color: white;
    }}
    .stDataFrame {{ padding-top: 10px; }}
</style>
""", unsafe_allow_html=True)


# ======================================================================================
# --- LÓGICA DE CARGA DE DATOS ---
# ======================================================================================

def normalizar_nombre(nombre: str) -> str:
    """Normaliza nombres de columnas y datos para comparación."""
    if not isinstance(nombre, str): return ""
    nombre = nombre.upper().strip().replace('.', '')
    nombre = ''.join(c for c in unicodedata.normalize('NFD', nombre) if unicodedata.category(c) != 'Mn')
    return ' '.join(nombre.split())

ZONAS_SERIE = { "PEREIRA": [155, 189, 158, 439], "MANIZALES": [157, 238], "ARMENIA": [156] }

def procesar_cartera(df: pd.DataFrame) -> pd.DataFrame:
    """Procesa el dataframe de cartera principal."""
    df_proc = df.copy()
    if 'importe' not in df_proc.columns: df_proc['importe'] = 0
    if 'numero' not in df_proc.columns: df_proc['numero'] = '0'
    if 'dias_vencido' not in df_proc.columns: df_proc['dias_vencido'] = 0
    if 'nomvendedor' not in df_proc.columns: df_proc['nomvendedor'] = None
    if 'serie' not in df_proc.columns: df_proc['serie'] = ''
    if 'fecha_documento' not in df_proc.columns: df_proc['fecha_documento'] = pd.NaT

    df_proc['importe'] = pd.to_numeric(df_proc['importe'], errors='coerce').fillna(0)
    df_proc['numero'] = df_proc['numero'].astype(str) 
    df_proc['serie'] = df_proc['serie'].astype(str) 
    df_proc['dias_vencido'] = pd.to_numeric(df_proc['dias_vencido'], errors='coerce').fillna(0)
    df_proc['fecha_documento'] = pd.to_datetime(df_proc['fecha_documento'], errors='coerce')
    df_proc['nomvendedor_norm'] = df_proc['nomvendedor'].apply(normalizar_nombre)
    
    ZONAS_SERIE_STR = {zona: [str(s) for s in series] for zona, series in ZONAS_SERIE.items()}
    
    def asignar_zona_robusta(valor_serie):
        if pd.isna(valor_serie): return "OTRAS ZONAS"
        numeros_en_celda = re.findall(r'\d+', str(valor_serie))
        if not numeros_en_celda: return "OTRAS ZONAS"
        for zona, series_clave_str in ZONAS_SERIE_STR.items():
            if set(numeros_en_celda) & set(series_clave_str): return zona
        return "OTRAS ZONAS"
    
    df_proc['zona'] = df_proc['serie'].apply(asignar_zona_robusta)
    bins = [-float('inf'), 0, 15, 30, 60, float('inf')]; labels = ['Al día', '1-15 días', '16-30 días', '31-60 días', 'Más de 60 días']
    df_proc['edad_cartera'] = pd.cut(df_proc['dias_vencido'], bins=bins, labels=labels, right=True)
    return df_proc

# --- Funciones de Carga de Dropbox ---

@st.cache_data(ttl=600)
def cargar_datos_cartera_dropbox():
    try:
        APP_KEY = st.secrets["dropbox"]["app_key"]
        APP_SECRET = st.secrets["dropbox"]["app_secret"]
        REFRESH_TOKEN = st.secrets["dropbox"]["refresh_token"]

        with dropbox.Dropbox(app_key=APP_KEY, app_secret=APP_SECRET, oauth2_refresh_token=REFRESH_TOKEN) as dbx:
            path_archivo_dropbox = '/data/cartera_detalle.csv'
            metadata, res = dbx.files_download(path=path_archivo_dropbox)
            contenido_csv = res.content.decode('latin-1')

            nombres_columnas_originales = [
                'Serie', 'Numero', 'Fecha Documento', 'Fecha Vencimiento', 'Cod Cliente',
                'NombreCliente', 'Nit', 'Poblacion', 'Provincia', 'Telefono1', 'Telefono2',
                'NomVendedor', 'Entidad Autoriza', 'E-Mail', 'Importe', 'Descuento',
                'Cupo Aprobado', 'Dias Vencido'
            ]

            df = pd.read_csv(
                StringIO(contenido_csv), 
                header=None, 
                names=nombres_columnas_originales, 
                sep='|', 
                engine='python',
                dtype={'Serie': str, 'Numero': str, 'Nit': str}
            )
            
            df_renamed = df.rename(columns=lambda x: normalizar_nombre(x).lower().replace(' ', '_'))
            df_renamed = df_renamed.loc[:, ~df_renamed.columns.duplicated()]
            df_renamed['fecha_documento'] = pd.to_datetime(df_renamed['fecha_documento'], errors='coerce')
            df_renamed['fecha_vencimiento'] = pd.to_datetime(df_renamed['fecha_vencimiento'], errors='coerce')
            
            return df_renamed
    except Exception as e:
        st.error(f"Error al cargar 'cartera_detalle.csv' desde Dropbox: {e}")
        return pd.DataFrame()

@st.cache_data(ttl=600)
def cargar_reporte_transacciones_dropbox():
    try:
        APP_KEY = st.secrets["dropbox"]["app_key"]
        APP_SECRET = st.secrets["dropbox"]["app_secret"]
        REFRESH_TOKEN = st.secrets["dropbox"]["refresh_token"]

        with dropbox.Dropbox(app_key=APP_KEY, app_secret=APP_SECRET, oauth2_refresh_token=REFRESH_TOKEN) as dbx:
            path_archivo_dropbox = '/data/reporteTransacciones.xlsx'
            metadata, res = dbx.files_download(path=path_archivo_dropbox)
            
            df = pd.read_excel(
                BytesIO(res.content),
                dtype={'DOCUMENTO': str, 'TITULO_VALOR': str, 'ESTADO': str}
            )
            
            df.columns = [normalizar_nombre(c).lower().replace(' ', '_') for c in df.columns]
            return df
    except Exception as e:
        st.error(f"Error al cargar 'reporteTransacciones.xlsx' desde Dropbox: {e}")
        return pd.DataFrame()

# --- Funciones de Normalización de Claves ---

def normalizar_nit_simple(nit_str: str) -> str:
    if not isinstance(nit_str, str): return ""
    return re.sub(r'\D', '', nit_str)

def obtener_nit_base(nit_str: str) -> str:
    nit_norm = normalizar_nit_simple(nit_str)
    if len(nit_norm) <= 1:
        return nit_norm
    return nit_norm[:-1]

def normalizar_factura_simple(fact_str: str) -> str:
    if not isinstance(fact_str, str): return ""
    return fact_str.split('.')[0].strip().upper().replace(' ', '').replace('-', '')

def normalizar_factura_cartera(row):
    serie = str(row['serie']).strip().upper()
    numero = str(row['numero']).split('.')[0].strip()
    return (serie + numero).replace(' ', '').replace('-', '')


# --- Función Principal de Procesamiento y Cruce ---

@st.cache_data
def cargar_y_comparar_datos():
    df_cartera_raw = cargar_datos_cartera_dropbox()
    if df_cartera_raw.empty:
        st.error("No se pudo cargar 'cartera_detalle.csv'.")
        # Retornamos los DataFrames vacíos para todas las variables
        return pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame()
    
    df_cartera = procesar_cartera(df_cartera_raw)
    
    # Filtro Series
    if 'serie' in df_cartera.columns:
        df_cartera = df_cartera[~df_cartera['serie'].astype(str).str.contains('W|X', case=False, na=False)]
        df_cartera = df_cartera[~df_cartera['serie'].astype(str).str.upper().str.endswith('U', na=False)]

    df_covinoc = cargar_reporte_transacciones_dropbox()
    if df_covinoc.empty:
        st.error("No se pudo cargar 'reporteTransacciones.xlsx'.")
        return pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame()

    df_cartera['nit_norm_cartera'] = df_cartera['nit'].apply(normalizar_nit_simple)
    set_nits_cartera = set(df_cartera['nit_norm_cartera'].unique())

    def encontrar_nit_en_cartera(doc_str_covinoc):
        if not isinstance(doc_str_covinoc, str): return None
        doc_norm = normalizar_nit_simple(doc_str_covinoc)
        if doc_norm in set_nits_cartera: return doc_norm
        doc_norm_base = doc_norm[:-1] 
        if doc_norm_base in set_nits_cartera: return doc_norm_base 
        return None 

    df_covinoc['nit_norm_cartera'] = df_covinoc['documento'].apply(encontrar_nit_en_cartera)
    df_cartera['factura_norm'] = df_cartera.apply(normalizar_factura_cartera, axis=1)
    df_covinoc['factura_norm'] = df_covinoc['titulo_valor'].apply(normalizar_factura_simple)

    df_cartera['clave_unica'] = df_cartera['nit_norm_cartera'] + '_' + df_cartera['factura_norm']
    df_covinoc['clave_unica'] = df_covinoc['nit_norm_cartera'] + '_' + df_covinoc['factura_norm']
    df_covinoc['estado_norm'] = df_covinoc['estado'].astype(str).str.upper().str.strip()
    
    # --- Tab 4 ---
    df_reclamadas = df_covinoc[df_covinoc['estado_norm'] == 'RECLAMADA'].copy()
    
    # --- Tab 1 ---
    nits_protegidos = df_covinoc['nit_norm_cartera'].dropna().unique()
    df_cartera_protegida = df_cartera[df_cartera['nit_norm_cartera'].isin(nits_protegidos)].copy()
    set_claves_covinoc_total = set(df_covinoc['clave_unica'].dropna().unique())
    df_a_subir_raw = df_cartera_protegida[~df_cartera_protegida['clave_unica'].isin(set_claves_covinoc_total)].copy()

    today = pd.to_datetime(datetime.now().date())
    if 'fecha_documento' in df_a_subir_raw.columns:
        df_a_subir_raw['dias_emision'] = (today - df_a_subir_raw['fecha_documento']).dt.days
        df_a_subir = df_a_subir_raw[(df_a_subir_raw['dias_emision'] >= 1) & (df_a_subir_raw['dias_emision'] <= 5)].copy()
    else:
        df_a_subir = df_a_subir_raw.iloc[0:0].copy() 

    # --- Tab 2 ---
    estados_cerrados = ['EFECTIVA', 'NEGADA', 'EXONERADA']
    df_covinoc_comparable = df_covinoc[~df_covinoc['estado_norm'].isin(estados_cerrados)].copy()
    set_claves_cartera_total = set(df_cartera['clave_unica'].dropna().unique())
    df_a_exonerar = df_covinoc_comparable[
        (~df_covinoc_comparable['clave_unica'].isin(set_claves_cartera_total)) &
        (df_covinoc_comparable['nit_norm_cartera'].notna())
    ].copy()

    # --- Intersección ---
    df_interseccion = pd.merge(df_cartera, df_covinoc, on='clave_unica', how='inner', suffixes=('_cartera', '_covinoc'))
    
    columnas_a_renombrar = {
        'importe': 'importe_cartera', 'nombrecliente': 'nombrecliente_cartera', 'nit': 'nit_cartera',
        'nomvendedor': 'nomvendedor_cartera', 'fecha_vencimiento': 'fecha_vencimiento_cartera',
        'dias_vencido': 'dias_vencido_cartera', 'saldo': 'saldo_covinoc', 'estado': 'estado_covinoc',
        'estado_norm': 'estado_norm_covinoc', 'vencimiento': 'vencimiento_covinoc'
    }
    cols_existentes = df_interseccion.columns
    renombres_aplicables = {k: v for k, v in columnas_a_renombrar.items() if k in cols_existentes}
    df_interseccion.rename(columns=renombres_aplicables, inplace=True)
    
    # --- Tab 3 ---
    df_aviso_no_pago_base = df_interseccion[df_interseccion['dias_vencido_cartera'] >= 25].copy()
    df_aviso_no_pago = df_aviso_no_pago_base[
        (pd.to_numeric(df_aviso_no_pago_base['importe_cartera'], errors='coerce').fillna(0) > 0) &
        (df_aviso_no_pago_base['estado_norm_covinoc'] != 'EXONERADA') &
        (df_aviso_no_pago_base['estado_norm_covinoc'] != 'NEGADA')
    ].copy()

    # --- Tab 5 ---
    df_interseccion['importe_cartera'] = pd.to_numeric(df_interseccion['importe_cartera'], errors='coerce').fillna(0)
    df_interseccion['saldo_covinoc'] = pd.to_numeric(df_interseccion['saldo_covinoc'], errors='coerce').fillna(0)
    df_ajustes = df_interseccion[(df_interseccion['saldo_covinoc'] > df_interseccion['importe_cartera'])].copy()
    df_ajustes['diferencia'] = df_ajustes['saldo_covinoc'] - df_ajustes['importe_cartera']

    # --- RETORNO AMPLIADO: INCLUYE LOS DATAFRAMES CRUDOS PARA REPORTES TOTALES ---
    return df_a_subir, df_a_exonerar, df_aviso_no_pago, df_reclamadas, df_ajustes, df_covinoc, df_cartera


# ======================================================================================
# --- FUNCIONES AUXILIARES PARA EXCEL ---
# ======================================================================================

def get_tipo_doc_from_nit_col(nit_str_raw: str) -> str:
    if not isinstance(nit_str_raw, str) or pd.isna(nit_str_raw): return 'C' 
    nit_str_raw_clean = nit_str_raw.strip().upper()
    if '-' in nit_str_raw_clean: return 'N'
    nit_norm = re.sub(r'\D', '', nit_str_raw_clean)
    if len(nit_norm) == 0: return 'C' 
    if (nit_norm.startswith('8') or nit_norm.startswith('9')): return 'N'
    return 'C'

def format_date(date_obj) -> str:
    if pd.isna(date_obj): return None
    try: return pd.to_datetime(date_obj).strftime('%Y/%m/%d')
    except Exception: return None

def to_excel(df: pd.DataFrame) -> bytes:
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Facturas')
    return output.getvalue()

def to_excel_informativo(df: pd.DataFrame) -> bytes:
    output = BytesIO()
    df_export = df.copy()
    mapa_columnas = {
        'nombrecliente': 'Cliente', 'nit': 'NIT', 'serie': 'Serie', 'numero': 'Factura',
        'factura_norm': 'Titulo Valor', 'fecha_documento': 'Fecha Emisión', 'dias_emision': 'Días desde Emisión',
        'fecha_vencimiento': 'Fecha Vencimiento', 'dias_vencido': 'Días Vencido',
        'importe': 'Valor Total', 'nomvendedor': 'Vendedor'
    }
    df_export = df_export.rename(columns=mapa_columnas)
    cols_deseadas = ['Cliente', 'NIT', 'Serie', 'Factura', 'Fecha Emisión', 'Días desde Emisión', 'Fecha Vencimiento', 'Días Vencido', 'Valor Total', 'Vendedor']
    cols_finales = [c for c in cols_deseadas if c in df_export.columns]
    df_export = df_export[cols_finales]
    if 'Días Vencido' in df_export.columns:
        df_export = df_export.sort_values(by='Días Vencido', ascending=False)
        
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        sheet_name = 'Reporte Detallado'
        df_export.to_excel(writer, index=False, sheet_name=sheet_name)
        workbook = writer.book
        worksheet = writer.sheets[sheet_name]
        # COLOR INSTITUCIONAL EN EXCEL (ROJO OSCURO)
        header_format = workbook.add_format({'bold': True, 'text_wrap': True, 'valign': 'top', 'fg_color': '#B21917', 'font_color': '#FFFFFF', 'border': 1})
        money_format = workbook.add_format({'num_format': '$ #,##0'})
        date_format = workbook.add_format({'num_format': 'yyyy-mm-dd'})
        
        for col_num, value in enumerate(df_export.columns.values):
            worksheet.write(0, col_num, value, header_format)
            if value == 'Cliente': worksheet.set_column(col_num, col_num, 40)
            elif value in ['NIT', 'Vendedor']: worksheet.set_column(col_num, col_num, 20)
            elif value in ['Valor Total']: worksheet.set_column(col_num, col_num, 18, money_format)
            elif 'Fecha' in value: worksheet.set_column(col_num, col_num, 15, date_format)
            else: worksheet.set_column(col_num, col_num, 15)
        
        max_row = len(df_export)
        worksheet.autofilter(0, 0, max_row, len(df_export.columns) - 1)
        if 'Días Vencido' in df_export.columns:
            idx_vencido = df_export.columns.get_loc('Días Vencido')
            col_letter = chr(ord('A') + idx_vencido) 
            rango_celdas = f"{col_letter}2:{col_letter}{max_row+1}"
            worksheet.conditional_format(rango_celdas, {'type': '3_color_scale', 'min_color': '#63BE7B', 'mid_color': '#FFEB84', 'max_color': '#F8696B'})
            
    return output.getvalue()

def to_excel_clientes_revision(df_resumen: pd.DataFrame) -> bytes:
    """
    Genera un Excel diseñado específicamente para que los compañeros revisen
    y marquen 'SI/NO' para depurar la base de datos.
    """
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        sheet_name = 'Revisión Clientes'
        # Renombrar columnas para el usuario final
        df_final = df_resumen.rename(columns={
            'nombrecliente': 'Cliente',
            'nit': 'NIT',
            'nomvendedor': 'Vendedor Principal',
            'importe': 'Deuda Total (Reporte)',
            'numero': 'Cantidad Facturas'
        })
        
        # Columnas vacías para diligenciar
        df_final['¿SEGUIR SUBIENDO? (SI/NO)'] = ''
        df_final['OBSERVACIONES'] = ''
        
        # Ordenar columnas
        cols = ['Cliente', 'NIT', 'Vendedor Principal', 'Cantidad Facturas', 'Deuda Total (Reporte)', '¿SEGUIR SUBIENDO? (SI/NO)', 'OBSERVACIONES']
        # Asegurar que existan, si no, crear vacías
        for c in cols:
            if c not in df_final.columns:
                df_final[c] = ''
                
        df_final = df_final[cols]
        
        df_final.to_excel(writer, index=False, sheet_name=sheet_name)
        
        workbook = writer.book
        worksheet = writer.sheets[sheet_name]
        
        # Formatos
        header_format = workbook.add_format({'bold': True, 'text_wrap': True, 'valign': 'top', 'fg_color': '#B21917', 'font_color': '#FFFFFF', 'border': 1})
        money_format = workbook.add_format({'num_format': '$ #,##0'})
        input_format = workbook.add_format({'bg_color': '#FEF4C0', 'border': 1}) # Amarillo suave para input
        
        for col_num, value in enumerate(df_final.columns.values):
            worksheet.write(0, col_num, value, header_format)
            if value == 'Cliente': worksheet.set_column(col_num, col_num, 45)
            elif value == 'NIT': worksheet.set_column(col_num, col_num, 15)
            elif value == 'Vendedor Principal': worksheet.set_column(col_num, col_num, 25)
            elif value == 'Deuda Total (Reporte)': worksheet.set_column(col_num, col_num, 20, money_format)
            elif value == '¿SEGUIR SUBIENDO? (SI/NO)': 
                worksheet.set_column(col_num, col_num, 25, input_format)
            elif value == 'OBSERVACIONES': 
                worksheet.set_column(col_num, col_num, 40, input_format)
            else:
                worksheet.set_column(col_num, col_num, 15)
                
        worksheet.autofilter(0, 0, len(df_final), len(df_final.columns) - 1)
        
    return output.getvalue()

def primer_valor_no_vacio(serie: pd.Series):
    valores = [valor for valor in serie if pd.notna(valor) and str(valor).strip()]
    return valores[0] if valores else ""

def unir_valores_unicos(serie: pd.Series) -> str:
    valores = sorted({str(valor).strip() for valor in serie if pd.notna(valor) and str(valor).strip()})
    return ' | '.join(valores)

def es_fau_digital_faltante(valor) -> bool:
    if pd.isna(valor):
        return True
    valor_norm = normalizar_nombre(str(valor)).replace(' ', '')
    return valor_norm in {'', 'NO', 'N', '0', 'SIN', 'NOAPLICA', 'NA', 'N/A', 'NULL', 'NONE', 'FALSE', 'PENDIENTE', 'FALTA'}

def leer_reporte_cupos_excel(origen_excel) -> pd.DataFrame:
    return pd.read_excel(
        origen_excel,
        dtype={
            'TIPO_DOCUMENTO': str,
            'DOCUMENTO': str,
            'FAU_DIGITAL': str,
            'PAGARE_DIGITAL': str,
            'USUARIO_SOLICITA': str,
            'USUARIO_GESTION': str,
            'SUCURSAL': str
        }
    )

@st.cache_data(ttl=600)
def cargar_reporte_cupos_local():
    rutas_encontradas = []
    directorios_busqueda = [
        os.getcwd(),
        os.path.dirname(os.path.abspath(__file__)),
        os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    ]
    patrones = ['reporteCupos*.xlsx', 'reporteCupos*.xls', 'ReporteCupos*.xlsx', 'ReporteCupos*.xls']

    for directorio in directorios_busqueda:
        for patron in patrones:
            rutas_encontradas.extend(glob.glob(os.path.join(directorio, patron)))

    rutas_unicas = []
    for ruta in rutas_encontradas:
        if ruta not in rutas_unicas:
            rutas_unicas.append(ruta)

    for ruta in rutas_unicas:
        try:
            return leer_reporte_cupos_excel(ruta), ruta, ""
        except Exception:
            continue

    return pd.DataFrame(), "", ""

@st.cache_data(ttl=600)
def cargar_reporte_cupos_dropbox():
    try:
        APP_KEY = st.secrets['dropbox']['app_key']
        APP_SECRET = st.secrets['dropbox']['app_secret']
        REFRESH_TOKEN = st.secrets['dropbox']['refresh_token']

        rutas_candidatas = [
            '/data/reporteCupos.xlsx',
            '/data/reportecupos.xlsx',
            '/reporteCupos.xlsx',
            '/data/reporteCupos.xls',
            '/reporteCupos.xls'
        ]

        with dropbox.Dropbox(app_key=APP_KEY, app_secret=APP_SECRET, oauth2_refresh_token=REFRESH_TOKEN) as dbx:
            for ruta in rutas_candidatas:
                try:
                    _, res = dbx.files_download(path=ruta)
                    return leer_reporte_cupos_excel(BytesIO(res.content)), ruta, ""
                except Exception:
                    continue

        return pd.DataFrame(), "", "No se encontró 'reporteCupos' en las rutas esperadas de Dropbox."
    except Exception as e:
        return pd.DataFrame(), "", f"Error al buscar 'reporteCupos' en Dropbox: {e}"

def obtener_reporte_cupos_df(uploaded_file=None):
    if uploaded_file is not None:
        try:
            return leer_reporte_cupos_excel(uploaded_file), 'Archivo cargado manualmente', ''
        except Exception as e:
            return pd.DataFrame(), '', f"No fue posible leer el archivo cargado: {e}"

    df_local, ruta_local, error_local = cargar_reporte_cupos_local()
    if not df_local.empty:
        return df_local, f'Archivo local: {ruta_local}', ''

    df_dropbox, ruta_dropbox, error_dropbox = cargar_reporte_cupos_dropbox()
    if not df_dropbox.empty:
        return df_dropbox, f'Dropbox: {ruta_dropbox}', ''

    mensaje = error_dropbox or error_local or "No se encontró el archivo 'reporteCupos'."
    return pd.DataFrame(), '', mensaje

def preparar_reporte_cupos(df_reporte_cupos: pd.DataFrame) -> pd.DataFrame:
    if df_reporte_cupos.empty:
        return pd.DataFrame()

    df = df_reporte_cupos.copy()
    df.columns = [normalizar_nombre(c).lower().replace(' ', '_') for c in df.columns]

    columnas_presentes = set(df.columns)
    columnas_esenciales = ['documento', 'nombres', 'fau_digital']
    columnas_faltantes = [col for col in columnas_esenciales if col not in columnas_presentes]
    if columnas_faltantes:
        raise ValueError(f"El archivo reporteCupos no contiene las columnas obligatorias: {', '.join(columnas_faltantes)}")

    columnas_base = [
        'tipo_documento', 'documento', 'nombres', 'estado', 'cupo_asignado', 'extracupo',
        'cupo_disponible', 'alerta', 'fecha_apertura', 'usuario_solicita', 'tipo_firma',
        'fau_digital', 'pagare_digital', 'usuario_gestion', 'sucursal'
    ]
    for columna in columnas_base:
        if columna not in df.columns:
            df[columna] = pd.NaT if columna == 'fecha_apertura' else ''

    columnas_texto = [
        'tipo_documento', 'documento', 'nombres', 'estado', 'alerta', 'usuario_solicita',
        'tipo_firma', 'fau_digital', 'pagare_digital', 'usuario_gestion', 'sucursal'
    ]
    for columna in columnas_texto:
        df[columna] = df[columna].fillna('').astype(str).str.strip()

    for columna in ['cupo_asignado', 'extracupo', 'cupo_disponible']:
        df[columna] = pd.to_numeric(df[columna], errors='coerce').fillna(0)

    df['fecha_apertura'] = pd.to_datetime(df['fecha_apertura'], errors='coerce')
    df['documento_norm'] = df['documento'].apply(normalizar_nit_simple)
    # Normalizar tipo_documento para mostrarlo y evitar errores de minúsculas/mayúsculas
    df['tipo_documento_norm'] = df['tipo_documento'].str.upper().str.strip().replace({'NIT':'N','CEDULA':'C','CÉDULA':'C'})
    df = df[df['documento_norm'] != ''].copy()
    df['fau_digital_faltante'] = df['fau_digital'].apply(es_fau_digital_faltante)
    # No filtrar por tipo_documento, incluir todos los tipos (C y N)
    return df

def construir_reporte_fau_pendiente(df_cartera_full: pd.DataFrame, df_reporte_cupos: pd.DataFrame):
    if df_cartera_full.empty or df_reporte_cupos.empty:
        return pd.DataFrame(), pd.DataFrame()

    df_cartera = df_cartera_full.copy()
    if 'nit_norm_cartera' not in df_cartera.columns:
        df_cartera['nit_norm_cartera'] = df_cartera['nit'].apply(normalizar_nit_simple)

    df_cartera['importe'] = pd.to_numeric(df_cartera['importe'], errors='coerce').fillna(0)
    df_cartera['dias_vencido'] = pd.to_numeric(df_cartera['dias_vencido'], errors='coerce').fillna(0)
    conteo_facturas_col = 'clave_unica' if 'clave_unica' in df_cartera.columns else 'numero'
    set_nits_cartera = set(df_cartera['nit_norm_cartera'].dropna().unique())

    def encontrar_nit_en_cartera_fau(doc_str_reporte):
        if not isinstance(doc_str_reporte, str):
            return None
        doc_norm = normalizar_nit_simple(doc_str_reporte)
        if doc_norm in set_nits_cartera:
            return doc_norm
        doc_norm_base = doc_norm[:-1]
        if doc_norm_base in set_nits_cartera:
            return doc_norm_base
        return None

    df_cartera_resumen = df_cartera.groupby('nit_norm_cartera').agg(
        nit=('nit', primer_valor_no_vacio),
        cliente=('nombrecliente', primer_valor_no_vacio),
        vendedor=('nomvendedor', primer_valor_no_vacio),
        facturas_activas=(conteo_facturas_col, 'nunique'),
        saldo_cartera=('importe', 'sum'),
        max_dias_vencido=('dias_vencido', 'max'),
        fecha_ultima_factura=('fecha_documento', 'max')
    ).reset_index()

    df_fau_faltante = df_reporte_cupos[df_reporte_cupos['fau_digital_faltante']].copy()
    if df_fau_faltante.empty:
        return pd.DataFrame(), pd.DataFrame()

    df_fau_faltante['tipo_documento_norm'] = df_fau_faltante['tipo_documento_norm'].replace('', 'S/D')
    df_fau_faltante['nit_norm_cartera'] = df_fau_faltante['documento'].apply(encontrar_nit_en_cartera_fau)

    df_fau_resumen = df_fau_faltante.groupby('documento_norm').agg(
        documento=('documento', primer_valor_no_vacio),
        tipo_documento=('tipo_documento_norm', primer_valor_no_vacio),
        nombres_reporte=('nombres', primer_valor_no_vacio),
        estado_cupo=('estado', unir_valores_unicos),
        tipo_firma=('tipo_firma', unir_valores_unicos),
        fau_digital=('fau_digital', unir_valores_unicos),
        pagare_digital=('pagare_digital', unir_valores_unicos),
        alerta=('alerta', unir_valores_unicos),
        sucursal=('sucursal', unir_valores_unicos),
        cupo_asignado=('cupo_asignado', 'max'),
        extracupo=('extracupo', 'max'),
        cupo_disponible=('cupo_disponible', 'max'),
        fecha_apertura=('fecha_apertura', 'max'),
        registros_reporte=('documento_norm', 'size'),
        nit_norm_cartera=('nit_norm_cartera', primer_valor_no_vacio)
    ).reset_index()

    df_consolidado = pd.merge(
        df_fau_resumen,
        df_cartera_resumen,
        left_on='nit_norm_cartera',
        right_on='nit_norm_cartera',
        how='left'
    )

    df_consolidado['cliente'] = df_consolidado['cliente'].fillna('').astype(str).str.strip()
    df_consolidado['nombres_reporte'] = df_consolidado['nombres_reporte'].fillna('').astype(str).str.strip()
    df_consolidado.loc[df_consolidado['cliente'] == '', 'cliente'] = df_consolidado['nombres_reporte']
    df_consolidado.loc[df_consolidado['cliente'] == '', 'cliente'] = 'CLIENTE SIN NOMBRE EN REPORTE'
    df_consolidado['vendedor'] = df_consolidado['vendedor'].fillna('').astype(str).str.strip()
    df_consolidado.loc[df_consolidado['vendedor'] == '', 'vendedor'] = 'GESTION INTERNA'
    df_consolidado['estado_cupo'] = df_consolidado['estado_cupo'].replace('', 'Sin estado reportado')
    df_consolidado['tipo_firma'] = df_consolidado['tipo_firma'].replace('', 'Sin tipo de firma')
    df_consolidado['alerta'] = df_consolidado['alerta'].replace('', 'Sin alerta')
    df_consolidado['fau_digital'] = 'PENDIENTE / VACIO'
    df_consolidado['vendedor_norm'] = df_consolidado['vendedor'].apply(normalizar_nombre)
    df_consolidado['relacion_vendedor'] = 'Cartera actual'
    df_consolidado.loc[df_consolidado['nit'].isna(), 'relacion_vendedor'] = 'GESTION INTERNA'
    df_consolidado['nit'] = df_consolidado['nit'].fillna('')
    df_consolidado['saldo_cartera'] = pd.to_numeric(df_consolidado['saldo_cartera'], errors='coerce').fillna(0)
    df_consolidado['facturas_activas'] = pd.to_numeric(df_consolidado['facturas_activas'], errors='coerce').fillna(0).astype(int)
    df_consolidado['max_dias_vencido'] = pd.to_numeric(df_consolidado['max_dias_vencido'], errors='coerce').fillna(0).astype(int)

    columnas_finales = [
        'vendedor', 'vendedor_norm', 'cliente', 'nit', 'documento', 'tipo_documento', 'estado_cupo',
        'tipo_firma', 'fau_digital', 'pagare_digital', 'cupo_asignado', 'extracupo', 'cupo_disponible',
        'saldo_cartera', 'facturas_activas', 'max_dias_vencido', 'fecha_ultima_factura', 'fecha_apertura',
        'sucursal', 'alerta', 'registros_reporte', 'relacion_vendedor'
    ]
    df_consolidado = df_consolidado[columnas_finales].copy()
    df_consolidado.rename(columns={
        'vendedor': 'Vendedor',
        'vendedor_norm': 'vendedor_norm',
        'cliente': 'Cliente',
        'nit': 'NIT Cartera',
        'documento': 'Documento Reporte',
        'tipo_documento': 'Tipo Documento',
        'estado_cupo': 'Estado Cupo',
        'tipo_firma': 'Tipo Firma',
        'fau_digital': 'FAU Digital',
        'pagare_digital': 'Pagare Digital',
        'cupo_asignado': 'Cupo Asignado',
        'extracupo': 'Extracupo',
        'cupo_disponible': 'Cupo Disponible',
        'saldo_cartera': 'Saldo Cartera',
        'facturas_activas': 'Facturas Activas',
        'max_dias_vencido': 'Max Dias Vencido',
        'fecha_ultima_factura': 'Fecha Ultima Factura',
        'fecha_apertura': 'Fecha Apertura Cupo',
        'sucursal': 'Sucursal',
        'alerta': 'Alerta',
        'registros_reporte': 'Registros Reporte',
        'relacion_vendedor': 'Relacion Vendedor'
    }, inplace=True)

    df_relacionados = df_consolidado[df_consolidado['Relacion Vendedor'] != 'GESTION INTERNA'].copy()
    df_no_relacionados = df_consolidado[df_consolidado['Relacion Vendedor'] == 'GESTION INTERNA'].copy()

    df_relacionados = df_relacionados.sort_values(by=['Vendedor', 'Cliente'], ascending=[True, True])
    df_no_relacionados = df_no_relacionados.sort_values(by=['Vendedor', 'Cliente'], ascending=[True, True])
    return df_relacionados, df_no_relacionados

def to_excel_fau_pendiente(df_fau: pd.DataFrame) -> bytes:
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        sheet_name = 'FAU Pendiente'
        df_fau.to_excel(writer, index=False, sheet_name=sheet_name)

        workbook = writer.book
        worksheet = writer.sheets[sheet_name]
        header_format = workbook.add_format({'bold': True, 'text_wrap': True, 'valign': 'top', 'fg_color': '#B21917', 'font_color': '#FFFFFF', 'border': 1})
        money_format = workbook.add_format({'num_format': '$ #,##0'})
        date_format = workbook.add_format({'num_format': 'yyyy-mm-dd'})

        for col_num, value in enumerate(df_fau.columns.values):
            worksheet.write(0, col_num, value, header_format)
            if value in ['Cliente', 'Sucursal', 'Alerta']:
                worksheet.set_column(col_num, col_num, 28)
            elif value in ['Vendedor', 'Estado Cupo', 'Tipo Firma']:
                worksheet.set_column(col_num, col_num, 24)
            elif value in ['Saldo Cartera', 'Cupo Asignado', 'Extracupo', 'Cupo Disponible']:
                worksheet.set_column(col_num, col_num, 18, money_format)
            elif value in ['Fecha Ultima Factura', 'Fecha Apertura Cupo']:
                worksheet.set_column(col_num, col_num, 18, date_format)
            else:
                worksheet.set_column(col_num, col_num, 18)

        worksheet.autofilter(0, 0, len(df_fau), len(df_fau.columns) - 1)

    return output.getvalue()


# ======================================================================================
# --- MÓDULO DE ANALÍTICA ESTRATÉGICA (Dashboard KPIs + Activación de Clientes) ---
#     Añadido sin alterar la lógica existente. Fuente: reporteTransacciones + reporteCupos
# ======================================================================================

BOLSA_COVINOC = 6_000_000_000                 # Bolsa total de garantía Covinoc: $6.000 millones
FECHA_INICIO_BOLSA = datetime(2025, 9, 30)    # Inicio del conteo de la bolsa (30-sep-2025)
ESTADOS_ACTIVOS_COVINOC = {'AL DIA', 'AVISO NO PAGO'}  # Títulos vigentes
PORTAL_PAGO = "https://ferreinoxtiendapintuco.epayco.me/recaudo/ferreinoxrecaudoenlinea/"
WHATSAPP_CARTERA = "573142087169"             # Línea de Cartera para activar cupo (+57 314 2087169)

MESES_ES = {1: 'Ene', 2: 'Feb', 3: 'Mar', 4: 'Abr', 5: 'May', 6: 'Jun',
            7: 'Jul', 8: 'Ago', 9: 'Sep', 10: 'Oct', 11: 'Nov', 12: 'Dic'}


def _etiqueta_periodo(periodo: str) -> str:
    """Convierte '2025-09' en 'Sep 2025'."""
    try:
        anio, mes = str(periodo).split('-')
        return f"{MESES_ES[int(mes)]} {anio}"
    except Exception:
        return str(periodo)


@st.cache_data(show_spinner=False)
def preparar_analitica_transacciones(df_covinoc: pd.DataFrame) -> pd.DataFrame:
    """Normaliza y enriquece el reporteTransacciones para el dashboard de KPIs."""
    if df_covinoc is None or df_covinoc.empty:
        return pd.DataFrame()

    df = df_covinoc.copy()
    df['fecha_dt'] = pd.to_datetime(df['fecha'], errors='coerce')
    df['vencimiento_dt'] = pd.to_datetime(df['vencimiento'], errors='coerce')
    df['valor_garantizado_num'] = pd.to_numeric(df['valor_garantizado'], errors='coerce').fillna(0)
    df['saldo_num'] = pd.to_numeric(df['saldo'], errors='coerce').fillna(0) if 'saldo' in df.columns else 0

    df['estado_norm'] = df['estado'].astype(str).str.upper().str.strip()
    df['es_activo'] = df['estado_norm'].isin(ESTADOS_ACTIVOS_COVINOC)

    if 'aviso_no_pago' in df.columns:
        df['tiene_aviso'] = df['aviso_no_pago'].apply(
            lambda v: pd.notna(v) and str(v).strip().lower() not in ('', 'none', 'nan')
        )
    else:
        df['tiene_aviso'] = False

    df['documento_norm'] = df['documento'].apply(normalizar_nit_simple) if 'documento' in df.columns else ''
    df['periodo'] = df['fecha_dt'].dt.to_period('M').astype(str)
    df.loc[df['fecha_dt'].isna(), 'periodo'] = pd.NA
    return df


def resumen_mensual_covinoc(dfa: pd.DataFrame) -> pd.DataFrame:
    """Serie mensual de valor garantizado desde el inicio de la bolsa."""
    if dfa is None or dfa.empty:
        return pd.DataFrame()
    d = dfa.dropna(subset=['fecha_dt']).copy()
    d = d[d['fecha_dt'] >= FECHA_INICIO_BOLSA]
    if d.empty:
        return pd.DataFrame()
    g = d.groupby('periodo').agg(
        valor=('valor_garantizado_num', 'sum'),
        titulos=('valor_garantizado_num', 'size'),
        clientes=('documento', 'nunique')
    ).reset_index().sort_values('periodo')
    g['acumulado'] = g['valor'].cumsum()
    g['mes_label'] = g['periodo'].apply(_etiqueta_periodo)
    return g


def analizar_cupos_clientes(df_cupos_prep: pd.DataFrame, dfa: pd.DataFrame) -> pd.DataFrame:
    """Cruza reporteCupos con transacciones para clasificar el uso del cupo por cliente."""
    if df_cupos_prep is None or df_cupos_prep.empty:
        return pd.DataFrame()

    dfc = df_cupos_prep.copy()
    if dfa is not None and not dfa.empty:
        act = dfa.groupby('documento_norm').agg(
            titulos_total=('valor_garantizado_num', 'size'),
            valor_total=('valor_garantizado_num', 'sum'),
            ultima_txn=('fecha_dt', 'max')
        ).reset_index()
    else:
        act = pd.DataFrame(columns=['documento_norm', 'titulos_total', 'valor_total', 'ultima_txn'])

    dfc = dfc.merge(act, on='documento_norm', how='left')
    dfc['titulos_total'] = pd.to_numeric(dfc['titulos_total'], errors='coerce').fillna(0).astype(int)
    dfc['valor_total'] = pd.to_numeric(dfc['valor_total'], errors='coerce').fillna(0)
    dfc['cupo_utilizado'] = (dfc['cupo_asignado'] - dfc['cupo_disponible']).clip(lower=0)
    dfc['tiene_cupo'] = dfc['cupo_asignado'] > 0
    dfc['usa_cupo'] = (dfc['cupo_utilizado'] > 0) | (dfc['titulos_total'] > 0)
    return dfc


def resumen_contacto_cartera(df_cartera_full: pd.DataFrame) -> pd.DataFrame:
    """Consolida datos de contacto (teléfono/correo) por NIT desde la cartera."""
    if df_cartera_full is None or df_cartera_full.empty:
        return pd.DataFrame()
    d = df_cartera_full.copy()
    if 'nit_norm_cartera' not in d.columns:
        d['nit_norm_cartera'] = d['nit'].apply(normalizar_nit_simple)
    for c in ['telefono1', 'email', 'nombrecliente', 'nomvendedor', 'cod_cliente']:
        if c not in d.columns:
            d[c] = ''
    g = d.groupby('nit_norm_cartera').agg(
        cliente_cartera=('nombrecliente', primer_valor_no_vacio),
        telefono=('telefono1', primer_valor_no_vacio),
        email=('email', primer_valor_no_vacio),
        vendedor_cartera=('nomvendedor', primer_valor_no_vacio),
        cod_cliente=('cod_cliente', primer_valor_no_vacio),
    ).reset_index()
    return g


def generar_link_wa_activacion(telefono, cliente, cupo_disponible) -> str:
    """Genera un link wa.me con mensaje de activación de cupo pre-cargado."""
    tel = re.sub(r'\D', '', str(telefono))
    if len(tel) == 10 and tel.startswith('3'):
        tel = '57' + tel
    if len(tel) < 10:
        return None
    nombre = str(cliente).split()[0].title() if cliente and str(cliente).strip() else "Cliente"
    try:
        cupo = float(cupo_disponible)
    except (TypeError, ValueError):
        cupo = 0
    cupo_txt = f"*${cupo:,.0f}*" if cupo > 0 else "un *cupo de crédito aprobado*"
    msg = (
        f"👋 Hola {nombre}, te saludamos de *Ferreinox SAS BIC*.\n\n"
        f"¡Tenemos buenas noticias! 🎉 Cuentas con {cupo_txt} de cupo de crédito, aprobado y listo para usar.\n\n"
        f"🛠️ Aprovéchalo en pinturas, ferretería y todo para tus proyectos... ¡y págalo después!\n\n"
        f"¿Te ayudamos a activarlo hoy? Responde este mensaje y con gusto te asesoramos. 🙌"
    )
    return f"https://wa.me/{tel}?text={urllib.parse.quote(msg)}"


def plantilla_activacion_html(cliente, cupo_disponible, vendedor) -> str:
    """Correo HTML institucional para campaña de activación de cupos."""
    nombre = str(cliente).split()[0].title() if cliente and str(cliente).strip() else "Cliente"
    try:
        cupo = float(cupo_disponible)
    except (TypeError, ValueError):
        cupo = 0
    cupo_bloque = f"${cupo:,.0f}" if cupo > 0 else "Cupo aprobado"
    vendedor_txt = str(vendedor).title() if vendedor and str(vendedor).strip() else "tu asesor Ferreinox"
    return f"""<!doctype html>
<html><head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1.0"></head>
<body style="margin:0;padding:0;background-color:#FAFAFA;font-family:'Segoe UI',Arial,sans-serif;">
  <table role="presentation" width="100%" cellpadding="0" cellspacing="0" style="background-color:#FAFAFA;padding:24px 0;">
    <tr><td align="center">
      <table role="presentation" width="600" cellpadding="0" cellspacing="0" style="max-width:600px;background:#FFFFFF;border-radius:16px;overflow:hidden;box-shadow:0 8px 24px rgba(0,0,0,0.08);">
        <tr><td style="background:linear-gradient(135deg,#B21917 0%,#E73537 60%,#F0833A 100%);padding:36px 32px;text-align:center;">
          <div style="color:#FFFFFF;font-size:13px;letter-spacing:3px;font-weight:600;opacity:.9;">FERREINOX SAS BIC</div>
          <div style="color:#FFFFFF;font-size:30px;font-weight:800;margin-top:8px;line-height:1.2;">Tu cupo de crédito<br>está listo 🎉</div>
        </td></tr>
        <tr><td style="padding:34px 40px 10px 40px;">
          <p style="font-size:17px;color:#31333F;margin:0 0 16px 0;">Hola <strong>{nombre}</strong>,</p>
          <p style="font-size:16px;color:#555;line-height:1.6;margin:0 0 24px 0;">
            En Ferreinox valoramos tu confianza. Por eso queremos recordarte que tienes un
            <strong>cupo de crédito aprobado</strong> y disponible para tus compras. ¡Actívalo y llévate hoy lo que necesitas, paga después!
          </p>
          <table role="presentation" width="100%" cellpadding="0" cellspacing="0" style="margin:8px 0 26px 0;">
            <tr><td style="background:#FEF4C0;border:2px dashed #F9B016;border-radius:14px;padding:22px;text-align:center;">
              <div style="font-size:13px;color:#B21917;font-weight:700;letter-spacing:1px;">CUPO DISPONIBLE</div>
              <div style="font-size:38px;color:#B21917;font-weight:800;margin-top:4px;">{cupo_bloque}</div>
              <div style="font-size:13px;color:#8a6d1a;margin-top:6px;">Listo para usar en tu próxima compra</div>
            </td></tr>
          </table>
          <table role="presentation" width="100%" cellpadding="0" cellspacing="0">
            <tr><td style="padding:6px 0;font-size:15px;color:#444;">🎨 &nbsp;Pinturas, esmaltes y acabados</td></tr>
            <tr><td style="padding:6px 0;font-size:15px;color:#444;">🛠️ &nbsp;Ferretería y herramientas</td></tr>
            <tr><td style="padding:6px 0;font-size:15px;color:#444;">🚚 &nbsp;Compra ahora y paga después, sin trámites</td></tr>
          </table>
          <div style="text-align:center;margin:30px 0 8px 0;">
            <a href="https://wa.me/{WHATSAPP_CARTERA}?text={urllib.parse.quote('Hola Cartera Ferreinox, quiero activar mi cupo de crédito. Mi empresa es: ' + str(cliente))}"
               style="background:#388E3C;color:#FFFFFF;text-decoration:none;font-size:16px;font-weight:700;padding:15px 40px;border-radius:30px;display:inline-block;">
               💬 Activar mi cupo con Cartera
            </a>
          </div>
          <p style="text-align:center;font-size:13px;color:#777;margin:14px 0 0 0;">
            O escríbenos a <strong>WhatsApp +57 314 2087169</strong> (línea de Cartera) y activamos tu cupo de inmediato.
          </p>
        </td></tr>
        <tr><td style="padding:18px 40px 34px 40px;border-top:1px solid #eee;text-align:center;">
          <p style="font-size:13px;color:#999;margin:14px 0 0 0;">Te atiende: <strong>{vendedor_txt}</strong> · Cartera Ferreinox</p>
          <p style="font-size:12px;color:#bbb;margin:6px 0 0 0;">Ferreinox SAS BIC · Este es un mensaje comercial de tu proveedor de confianza.</p>
        </td></tr>
      </table>
    </td></tr>
  </table>
</body></html>"""


def enviar_correo_activacion_sendgrid(api_key, from_email, from_name, to_email,
                                      cliente_nombre, subject, html_content, plain_content):
    """Envía un correo de activación vía SendGrid (mismo motor usado en la app)."""
    payload = {
        "personalizations": [{"to": [{"email": to_email, "name": str(cliente_nombre)}]}],
        "from": {"email": from_email, "name": from_name},
        "subject": subject,
        "content": [
            {"type": "text/plain", "value": plain_content},
            {"type": "text/html", "value": html_content},
        ],
        "tracking_settings": {
            "open_tracking": {"enable": True},
            "click_tracking": {"enable": True, "enable_text": False},
        },
        "categories": ["activacion-cupo", "covinoc", "ferreinox"],
    }
    request = urllib_request.Request(
        url="https://api.sendgrid.com/v3/mail/send",
        data=json.dumps(payload).encode("utf-8"),
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib_request.urlopen(request, timeout=45) as response:
            status = response.getcode()
            if 200 <= status < 300:
                return True, f"HTTP {status}"
            return False, f"HTTP {status}"
    except Exception as e:
        return False, str(e)


def to_excel_generico(hojas: dict) -> bytes:
    """Exporta múltiples DataFrames a un Excel con estética institucional."""
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        workbook = writer.book
        header_format = workbook.add_format({'bold': True, 'text_wrap': True, 'valign': 'top',
                                             'fg_color': '#B21917', 'font_color': '#FFFFFF', 'border': 1})
        money_format = workbook.add_format({'num_format': '$ #,##0'})
        for nombre_hoja, df_hoja in hojas.items():
            if df_hoja is None or df_hoja.empty:
                continue
            df_hoja.to_excel(writer, index=False, sheet_name=nombre_hoja[:31])
            worksheet = writer.sheets[nombre_hoja[:31]]
            for col_num, value in enumerate(df_hoja.columns.values):
                worksheet.write(0, col_num, value, header_format)
                ancho = 22
                if any(k in str(value).lower() for k in ['cliente', 'nombre', 'correo', 'email']):
                    ancho = 38
                fmt = money_format if any(k in str(value).lower() for k in ['valor', 'cupo', 'saldo', 'monto', 'acumulado']) else None
                worksheet.set_column(col_num, col_num, ancho, fmt)
            worksheet.autofilter(0, 0, len(df_hoja), len(df_hoja.columns) - 1)
    return output.getvalue()


# ======================================================================================
# --- NUEVA LÓGICA: GENERACIÓN DE DOCUMENTOS WORD PROFESIONALES (MEJORADO) ---
# ======================================================================================

def set_cell_background(cell, color_hex):
    """Establece el color de fondo de una celda de tabla."""
    tcPr = cell._element.tcPr
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex.replace('#', ''))
    tcPr.append(shd)

def aplicar_estilo_parrafo(p, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, color=None):
    """Aplica formato consistente Quicksand a un párrafo."""
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.runs[0] if p.runs else p.add_run()
    run.font.name = 'Quicksand' # CAMBIO DE FUENTE
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def crear_encabezado_profesional(doc, titulo_principal=None):
    """Crea un encabezado visualmente limpio."""
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    
    # Encabezado "Empresarial" (Texto)
    p = doc.add_paragraph()
    run = p.add_run("FERREINOX S.A.S. BIC.")
    run.font.name = 'Quicksand' # CAMBIO DE FUENTE
    run.font.size = Pt(18)
    run.font.bold = True
    # Rojo Institucional B21917 (RGB: 178, 25, 23)
    run.font.color.rgb = RGBColor(178, 25, 23) 
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    
    p2 = doc.add_paragraph("NIT: 800.224.617-8")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.runs[0] if p2.runs else p2.add_run()
    run2.font.name = 'Quicksand' # CAMBIO DE FUENTE
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(100, 100, 100) # Gris
    p2.paragraph_format.space_after = Pt(20) # Espacio antes del título del doc

    if titulo_principal:
        p_tit = doc.add_paragraph(titulo_principal)
        p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_tit = p_tit.runs[0] if p_tit.runs else p_tit.add_run()
        run_tit.font.name = 'Quicksand' # CAMBIO DE FUENTE
        run_tit.font.size = Pt(14)
        run_tit.font.bold = True
        run_tit.font.underline = True
        run_tit.font.color.rgb = RGBColor(178, 25, 23) # Rojo
        p_tit.paragraph_format.space_after = Pt(24)

def generar_documentos_reclamacion(cliente_nombre, cliente_nit, cliente_dir, cliente_ciudad, facturas_data):
    """
    Genera 3 documentos ZIP optimizados visualmente.
    Usa Quicksand, Tablas Amarillas Pálidas, Espaciados Correctos.
    """
    zip_buffer = BytesIO()
    meses = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
    f_now = datetime.now()
    fecha_larga = f"{f_now.day} de {meses[f_now.month-1]} de {f_now.year}"
    fecha_ciudad = f"Pereira, {fecha_larga}"

    # COLORES INSTITUCIONALES PARA WORD
    rojo_institucional = RGBColor(178, 25, 23) # #B21917
    fondo_tabla_header = "B21917" # Rojo Oscuro
    fondo_tabla_body = "FEF4C0"   # Amarillo Pálido Institucional

    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        
        # --- 1. NOTIFICACIÓN DEUDOR (Mejorada) ---
        doc = Document()
        crear_encabezado_profesional(doc, titulo_principal=None) # Sin título centrado, es carta

        # Fecha y Datos Destinatario
        p = doc.add_paragraph(fecha_ciudad)
        aplicar_estilo_parrafo(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=18)

        p = doc.add_paragraph("Señor(a):")
        aplicar_estilo_parrafo(p, bold=True, space_after=2)
        p = doc.add_paragraph(f"{cliente_nombre}")
        aplicar_estilo_parrafo(p, space_after=2)
        if cliente_dir and cliente_dir != "Sin Dirección":
            p = doc.add_paragraph(f"{cliente_dir}")
            aplicar_estilo_parrafo(p, space_after=2)
        p = doc.add_paragraph(f"{cliente_ciudad if cliente_ciudad else 'Ciudad'}")
        aplicar_estilo_parrafo(p, space_after=18)

        # Asunto
        p = doc.add_paragraph()
        run = p.add_run("REF: NOTIFICACIÓN DE ENDOSO DE TÍTULOS VALORES")
        run.font.bold = True
        run.font.name = 'Quicksand'
        run.font.size = Pt(11)
        run.font.color.rgb = rojo_institucional
        p.paragraph_format.space_after = Pt(18)

        # Cuerpo
        p = doc.add_paragraph("Respetado Señor(a):")
        aplicar_estilo_parrafo(p, space_after=12)
        
        texto_intro = (
            "Por medio de la presente queremos comunicarle que los siguientes Títulos Valores "
            "han sido endosados en propiedad a favor de NEGOCIACIÓN DE TÍTULOS NET S.A.S:"
        )
        p = doc.add_paragraph(texto_intro)
        aplicar_estilo_parrafo(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12)

        # Tabla Estilizada
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        
        # Anchos relativos
        table.columns[0].width = Inches(1.5) # Titulo
        table.columns[1].width = Inches(1.5) # Valor I
        table.columns[2].width = Inches(1.5) # Abono
        table.columns[3].width = Inches(1.5) # Final

        # Encabezados Tabla
        hdr_cells = table.rows[0].cells
        titulos = ['Título Valor', 'Valor Inicial', 'Abono', 'Valor Final']
        for i, t in enumerate(titulos):
            hdr_cells[i].text = t
            set_cell_background(hdr_cells[i], fondo_tabla_header)
            p_cell = hdr_cells[i].paragraphs[0]
            p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_cell = p_cell.runs[0]
            run_cell.font.bold = True
            run_cell.font.name = 'Quicksand'
            run_cell.font.size = Pt(10)
            run_cell.font.color.rgb = RGBColor(255, 255, 255) # Blanco

        # Datos Tabla
        total_deuda = 0
        for fac in facturas_data:
            row_cells = table.add_row().cells
            # Factura (Centro)
            row_cells[0].text = str(fac['factura'])
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            val = fac['valor']
            total_deuda += val
            val_fmt = f"${val:,.0f}"

            # Valores (Derecha)
            row_cells[1].text = val_fmt
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            row_cells[2].text = "$0"
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            row_cells[3].text = val_fmt
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Fuente tabla cuerpo y fondo amarillo pálido
            for c in row_cells:
                set_cell_background(c, fondo_tabla_body)
                c.paragraphs[0].runs[0].font.name = 'Quicksand'
                c.paragraphs[0].runs[0].font.size = Pt(10)

        doc.add_paragraph().paragraph_format.space_after = Pt(12)

        texto_legal = (
            "Por lo anterior, sus pagos a partir de la fecha deberán realizarse a favor de NEGOCIACIÓN DE TÍTULOS NET S.A.S. "
            "Es importante mencionarle que si sus obligaciones se encuentran al día, contará con los beneficios de mantener su "
            "buen comportamiento de pago y mantener su cupo activo de compra.\n"
            "Finalmente, le informamos que COVINOC como administrador de la cartera de NEGOCIACIÓN DE TÍTULOS NET S.A.S., "
            "atenderá cualquier inquietud relacionada con sus obligaciones. Agradecemos solicitar su orden de pago y proceder a cancelar su obligación."
        )
        p = doc.add_paragraph(texto_legal)
        aplicar_estilo_parrafo(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12)

        texto_contacto = (
            "Para mayor información puede comunicarse en Bogotá llamando a los teléfonos 3534311 o al 3534324, a nivel nacional 018000946969, "
            "o también puede escribir al correo electrónico cobranza.sep@covinoc.com."
        )
        p = doc.add_paragraph(texto_contacto)
        aplicar_estilo_parrafo(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=30)

        # Firma
        p = doc.add_paragraph("Cordialmente,")
        aplicar_estilo_parrafo(p, space_after=40)

        p = doc.add_paragraph("__________________________________________")
        p.paragraph_format.space_after = Pt(2)
        p = doc.add_paragraph("FERREINOX S.A.S. BIC.")
        aplicar_estilo_parrafo(p, bold=True, space_after=2, color=rojo_institucional)
        p = doc.add_paragraph("NIT: 800.224.617-8")
        aplicar_estilo_parrafo(p, size=10)

        # Guardar
        bio = BytesIO()
        doc.save(bio)
        zip_file.writestr(f"Notificacion_{normalizar_nit_simple(str(cliente_nit))}.docx", bio.getvalue())


        # --- 2. DOCUMENTO ENDOSO (Mejorado) ---
        doc = Document()
        crear_encabezado_profesional(doc, titulo_principal="ENDOSO EN PROPIEDAD")

        texto_endoso = (
            f"Yo, JORGE IVAN PEREZ ANGEL, mayor de edad, identificado como consta al pie de mi firma, "
            f"actuando en mi calidad de representante legal de FERREINOX S.A.S. BIC, identificada con el NIT 800.224.617-8, "
            f"manifiesto que ENDOSO EN PROPIEDAD a la orden de NEGOCIACIÓN DE TÍTULOS NET S.A.S., identificada con NIT 830.051.527-9, "
            f"las siguientes facturas de venta:"
        )
        p = doc.add_paragraph(texto_endoso)
        aplicar_estilo_parrafo(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=18)

        # Tabla Endoso
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        titulos_e = ['No. Factura', 'Fecha Vencimiento', 'Valor']
        hdr = table.rows[0].cells
        for i, t in enumerate(titulos_e):
            hdr[i].text = t
            set_cell_background(hdr[i], fondo_tabla_header)
            p_h = hdr[i].paragraphs[0]
            p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_h = p_h.runs[0]
            run_h.font.bold = True
            run_h.font.name = 'Quicksand'
            run_h.font.size = Pt(11)
            run_h.font.color.rgb = RGBColor(255, 255, 255)

        for fac in facturas_data:
            row = table.add_row().cells
            row[0].text = str(fac['factura'])
            row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            row[1].text = str(fac['fecha_venc'])
            row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            row[2].text = f"${fac['valor']:,.0f}"
            row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            for c in row:
                set_cell_background(c, fondo_tabla_body)
                c.paragraphs[0].runs[0].font.name = 'Quicksand'
                c.paragraphs[0].runs[0].font.size = Pt(11)

        doc.add_paragraph().paragraph_format.space_after = Pt(24)

        p = doc.add_paragraph(f"Para constancia se firma en la ciudad de Pereira, el día {fecha_larga}.")
        aplicar_estilo_parrafo(p, size=12, space_after=50)

        # Firma
        p = doc.add_paragraph("__________________________________________")
        p.paragraph_format.space_after = Pt(2)
        p = doc.add_paragraph("FIRMA DEL REPRESENTANTE LEGAL")
        aplicar_estilo_parrafo(p, bold=True, space_after=2, color=rojo_institucional)
        p = doc.add_paragraph("C.C. _______________________ de ________________")
        aplicar_estilo_parrafo(p, size=11)

        bio = BytesIO()
        doc.save(bio)
        zip_file.writestr(f"Endoso_{normalizar_nit_simple(str(cliente_nit))}.docx", bio.getvalue())


        # --- 3. ACEPTACIÓN TÁCITA (Mejorada) ---
        doc = Document()
        crear_encabezado_profesional(doc, titulo_principal="CONSTANCIA DE ACEPTACIÓN TÁCITA")

        p = doc.add_paragraph(fecha_ciudad)
        aplicar_estilo_parrafo(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=24)

        texto_tacita = (
            "Bajo la gravedad de juramento, me permito indicar que las facturas de venta relacionadas a continuación "
            "no han sido aceptadas expresamente; en tal sentido, han sido aceptadas tácitamente y no se ha efectuado "
            "reclamo o devolución de las mismas de acuerdo a lo estipulado en el Código de Comercio."
        )
        p = doc.add_paragraph(texto_tacita)
        aplicar_estilo_parrafo(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=18)

        # Tabla Tácita
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Encabezado
        hdr = table.rows[0].cells
        hdr[0].text = "Factura No."
        hdr[1].text = "Valor Total"
        for c in hdr:
            set_cell_background(c, fondo_tabla_header)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraphs[0].runs[0].font.bold = True
            c.paragraphs[0].runs[0].font.name = 'Quicksand'
            c.paragraphs[0].runs[0].font.size = Pt(11)
            c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        for fac in facturas_data:
            row = table.add_row().cells
            row[0].text = str(fac['factura'])
            row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            row[1].text = f"${fac['valor']:,.0f}"
            row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            for c in row:
                set_cell_background(c, fondo_tabla_body)
                c.paragraphs[0].runs[0].font.name = 'Quicksand'
                c.paragraphs[0].runs[0].font.size = Pt(11)

        doc.add_paragraph().paragraph_format.space_after = Pt(40)

        # Firma
        p = doc.add_paragraph("__________________________________________")
        p.paragraph_format.space_after = Pt(2)
        p = doc.add_paragraph("FERREINOX S.A.S. BIC.")
        aplicar_estilo_parrafo(p, bold=True, space_after=2, color=rojo_institucional)
        p = doc.add_paragraph("REPRESENTANTE LEGAL")
        aplicar_estilo_parrafo(p, size=11)

        bio = BytesIO()
        doc.save(bio)
        zip_file.writestr(f"Aceptacion_Tacita_{normalizar_nit_simple(str(cliente_nit))}.docx", bio.getvalue())

    return zip_buffer


# ======================================================================================
# --- BLOQUE PRINCIPAL DE LA APP ---
# ======================================================================================
def main():
    # --- Lógica de Autenticación ---
    if 'authentication_status' not in st.session_state:
        st.session_state['authentication_status'] = False
        st.session_state['acceso_general'] = False
        st.session_state['vendedor_autenticado'] = None

    if not st.session_state['authentication_status']:
        st.title("Acceso al Módulo de Cartera Protegida")
        try:
            general_password = st.secrets["general"]["password"]
            vendedores_secrets = st.secrets["vendedores"]
        except Exception as e:
            st.error(f"Error al cargar las contraseñas desde los secretos: {e}")
            st.stop()
        
        password = st.text_input("Introduce la contraseña:", type="password", key="password_input_covinoc")
        
        if st.button("Ingresar"):
            if password == str(general_password):
                st.session_state['authentication_status'] = True
                st.session_state['acceso_general'] = True
                st.session_state['vendedor_autenticado'] = "General"
                st.rerun()
            else:
                for vendedor_key, pass_vendedor in vendedores_secrets.items():
                    if password == str(pass_vendedor):
                        st.session_state['authentication_status'] = True
                        st.session_state['acceso_general'] = False
                        st.session_state['vendedor_autenticado'] = vendedor_key
                        st.rerun()
                        break
                if not st.session_state['authentication_status']:
                    st.error("Contraseña incorrecta.")
    else:
        # --- Aplicación Principal (Usuario Autenticado) ---
        st.title("🛡️ Gestión de Cartera Protegida (Covinoc)")

        if st.button("🔄 Recargar Datos (Dropbox)"):
            st.cache_data.clear()
            st.success("Caché limpiado. Recargando datos de Cartera y Covinoc...")
            st.rerun()

        # --- Barra Lateral (Sidebar) ---
        with st.sidebar:
            try:
                st.image("LOGO FERREINOX SAS BIC 2024.png", use_container_width=True)
            except Exception:
                st.warning("Logo no encontrado.")
            
            st.success(f"Usuario: {st.session_state['vendedor_autenticado']}")
            
            if st.button("Cerrar Sesión"):
                for key in list(st.session_state.keys()):
                    del st.session_state[key]
                st.rerun()
            
            st.markdown("---")
            st.info("Esta página compara la cartera de Ferreinox con el reporte de transacciones de Covinoc.")

        # --- Carga y Procesamiento de Datos ---
        with st.spinner("Cargando y comparando archivos de Dropbox..."):
            # AQUI SE DESEMPAQUETAN LOS 7 ELEMENTOS
            df_a_subir, df_a_exonerar, df_aviso_no_pago, df_reclamadas, df_ajustes, df_covinoc_full, df_cartera_full = cargar_y_comparar_datos()

        if df_a_subir.empty and df_a_exonerar.empty and df_aviso_no_pago.empty and df_reclamadas.empty and df_ajustes.empty:
            st.warning("Se cargaron los archivos, pero no se encontraron diferencias para las 5 categorías.")
            st.info("Nota: En la Pestaña 1, solo se muestran facturas con 1 a 5 días de emisión.")
        
        # --- Datos para el módulo de Analítica Estratégica (Tabs 7 y 8) ---
        dfa_covinoc = preparar_analitica_transacciones(df_covinoc_full)
        try:
            df_cupos_auto_raw, fuente_cupos_auto, error_cupos_auto = obtener_reporte_cupos_df()
            df_cupos_auto = preparar_reporte_cupos(df_cupos_auto_raw) if not df_cupos_auto_raw.empty else pd.DataFrame()
        except Exception:
            df_cupos_auto, fuente_cupos_auto, error_cupos_auto = pd.DataFrame(), '', ''

        # --- Contenedor Principal con Pestañas ---
        st.markdown("---")

        tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8 = st.tabs([
            f"1. Facturas a Subir ({len(df_a_subir)})",
            f"2. Exoneraciones ({len(df_a_exonerar)})",
            f"3. Avisos de No Pago ({len(df_aviso_no_pago)})",
            f"4. Reclamadas ({len(df_reclamadas)})",
            f"5. Ajustes Parciales ({len(df_ajustes)})",
            "6. FAU Digital Pendiente",
            "📊 7. Dashboard KPIs",
            "🚀 8. Activación Clientes"
        ])

        with tab1:
            st.subheader("Facturas a Subir a Covinoc")
            st.markdown("Facturas de **clientes protegidos** que están en **Cartera Ferreinox** pero **NO** en Covinoc.")
            st.warning("🚩 **Importante:** Esta lista ya está pre-filtrada para mostrar **ÚNICAMENTE** facturas con 1 a 5 días desde su fecha de emisión.")
            
            if df_a_subir.empty:
                st.info("No hay facturas pendientes por subir que cumplan el criterio de 1 a 5 días de emisión.")
            else:
                st.markdown("---")
                # =================================================================================
                # --- NUEVO: BOTÓN DE DESCARGA LISTADO COMPLETO COVINOC (SIN FILTROS) ---
                # =================================================================================
                st.markdown("##### 🛠️ Herramientas Administrativas")
                col_admin_1, col_admin_2 = st.columns([0.7, 0.3])
                with col_admin_1:
                    st.info("¿Desea descargar el listado TOTAL de clientes que existen en ReporteTransacciones (Covinoc)?")
                with col_admin_2:
                    # LÓGICA MODIFICADA PARA LISTAR TODOS LOS CLIENTES DE COVINOC (ReporteTransacciones)
                    # Y CRUZARLOS CON CARTERA PARA OBTENER EL NOMBRE Y VENDEDOR
                    if not df_covinoc_full.empty:
                        # 1. Agrupar la data de Covinoc (que es la fuente de verdad para este reporte)
                        df_covinoc_full['saldo'] = pd.to_numeric(df_covinoc_full['saldo'], errors='coerce').fillna(0)
                        
                        # Usamos 'nit_norm_cartera' que ya fue calculado en cargar_y_comparar_datos
                        # Si es nulo (no encontró match), usamos el documento original limpio
                        df_covinoc_full['nit_join'] = df_covinoc_full['nit_norm_cartera']
                        mask_sin_nit = df_covinoc_full['nit_join'].isna()
                        if mask_sin_nit.any():
                             df_covinoc_full.loc[mask_sin_nit, 'nit_join'] = df_covinoc_full.loc[mask_sin_nit, 'documento'].apply(normalizar_nit_simple)
                        
                        df_resumen_covinoc = df_covinoc_full.groupby('nit_join').agg({
                            'saldo': 'sum',
                            'titulo_valor': 'count'
                        }).reset_index()
                        
                        # 2. Obtener nombres y vendedores de Cartera Completa (únicos por NIT)
                        if not df_cartera_full.empty:
                            df_info_clientes = df_cartera_full[['nit_norm_cartera', 'nombrecliente', 'nomvendedor']].drop_duplicates(subset=['nit_norm_cartera'])
                            
                            # 3. Cruzar (Left Join) para traer info descriptiva a los clientes de Covinoc
                            df_final_revision = pd.merge(df_resumen_covinoc, df_info_clientes, left_on='nit_join', right_on='nit_norm_cartera', how='left')
                        else:
                            df_final_revision = df_resumen_covinoc.copy()
                            df_final_revision['nombrecliente'] = None
                            df_final_revision['nomvendedor'] = None
                        
                        # 4. Rellenar vacíos para clientes que están en Covinoc pero YA NO en Cartera actual
                        df_final_revision['nombrecliente'] = df_final_revision['nombrecliente'].fillna('CLIENTE EN COVINOC - NO EN CARTERA ACTUAL')
                        df_final_revision['nomvendedor'] = df_final_revision['nomvendedor'].fillna('S/N')
                        
                        # 5. Renombrar columnas para que coincida con la función exportadora
                        df_final_revision.rename(columns={
                            'saldo': 'importe',
                            'titulo_valor': 'numero',
                            'nit_join': 'nit'
                        }, inplace=True)
                        
                        # Generar Excel
                        excel_clientes_revision = to_excel_clientes_revision(df_final_revision)
                        
                        st.download_button(
                            label="📂 Descargar Listado TOTAL Clientes Covinoc",
                            data=excel_clientes_revision,
                            file_name="Listado_TOTAL_Clientes_Covinoc.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
                    else:
                        st.warning("No hay datos en reporteTransacciones para generar el listado.")

                st.markdown("---")

                st.subheader("Filtros Adicionales")
                
                # Filtro 1: Excluir Clientes
                clientes_unicos = sorted(df_a_subir['nombrecliente'].dropna().unique())
                clientes_excluidos = st.multiselect(
                    "1. Excluir Clientes del Listado:",
                    options=clientes_unicos,
                    default=[],
                    help="Seleccione uno o más clientes para ocultar sus facturas de la lista de selección."
                )
                
                # Filtro 2: Incluir Series
                series_options_base = ['156', '157', '158', '189', '238', '439']
                series_disponibles_en_df = sorted(df_a_subir['serie'].dropna().astype(str).unique())
                series_options_final = sorted(list(set(series_options_base + series_disponibles_en_df)))
                
                series_seleccionadas = st.multiselect(
                    "2. Filtrar por Serie (Seleccione una o varias):",
                    options=series_options_final,
                    default=series_disponibles_en_df, 
                    help="Seleccione las series de factura que desea incluir en la lista."
                )

                # Filtro 3: Días Vencido
                if not df_a_subir['dias_vencido'].empty:
                    min_dias = int(df_a_subir['dias_vencido'].min())
                    max_dias = int(df_a_subir['dias_vencido'].max())
                    if min_dias == max_dias: max_dias += 1
                else:
                    min_dias, max_dias = 0, 1
                    
                dias_range = st.slider(
                    "3. Filtrar por Días Vencido:", 
                    min_value=min_dias, 
                    max_value=max_dias, 
                    value=(min_dias, max_dias),
                    help="Seleccione el rango de días de vencimiento a incluir."
                )
                
                # Aplicar TODOS los filtros
                df_a_subir_filtrado = df_a_subir[
                    (~df_a_subir['nombrecliente'].isin(clientes_excluidos)) &
                    (df_a_subir['serie'].astype(str).isin(series_seleccionadas)) &
                    (df_a_subir['dias_vencido'] >= dias_range[0]) &
                    (df_a_subir['dias_vencido'] <= dias_range[1])
                ].copy()

                st.markdown("---")
                st.subheader("Indicadores de Gestión (Facturas Filtradas)")
                
                kpi_col1, kpi_col2, kpi_col3 = st.columns(3)
                try:
                    monto_total_filtrado = pd.to_numeric(df_a_subir_filtrado['importe'], errors='coerce').sum()
                    clientes_unicos_filtrados = df_a_subir_filtrado['nombrecliente'].nunique()
                except Exception:
                    monto_total_filtrado = 0
                    clientes_unicos_filtrados = 0

                kpi_col1.metric("Nº Facturas (Filtradas)", f"{len(df_a_subir_filtrado)}")
                kpi_col2.metric("Monto Total (Filtrado)", f"${monto_total_filtrado:,.0f}")
                kpi_col3.metric("Nº Clientes (Filtrados)", f"{clientes_unicos_filtrados}")
                st.markdown("---")
                
                st.subheader("Selección de Facturas para Descarga")
                st.info("Utilice las casillas de la columna 'Seleccionar' para elegir qué facturas desea incluir en el archivo Excel.")

                # Lógica de Botones "Seleccionar Todos"
                if 'data_editor_key_tab1' not in st.session_state:
                    st.session_state.data_editor_key_tab1 = "data_editor_subir_0"
                if 'default_select_val_tab1' not in st.session_state:
                    st.session_state.default_select_val_tab1 = False

                sel_col1, sel_col2 = st.columns(2)
                with sel_col1:
                    if st.button("✅ Seleccionar Todos (Visible en Filtro)", use_container_width=True):
                        st.session_state.default_select_val_tab1 = True
                        st.session_state.data_editor_key_tab1 = f"data_editor_subir_{int(st.session_state.data_editor_key_tab1.split('_')[-1]) + 1}"
                        st.rerun()
                with sel_col2:
                    if st.button("◻️ Deseleccionar Todos (Visible en Filtro)", use_container_width=True):
                        st.session_state.default_select_val_tab1 = False
                        st.session_state.data_editor_key_tab1 = f"data_editor_subir_{int(st.session_state.data_editor_key_tab1.split('_')[-1]) + 1}"
                        st.rerun()

                columnas_mostrar_subir = [
                    'nombrecliente', 'nit', 'serie', 'numero', 'factura_norm', 
                    'fecha_documento', 'dias_emision',
                    'fecha_vencimiento', 'dias_vencido', 'importe', 'nomvendedor', 'clave_unica'
                ]
                columnas_existentes_subir = [col for col in columnas_mostrar_subir if col in df_a_subir_filtrado.columns]
                
                df_para_seleccionar = df_a_subir_filtrado[columnas_existentes_subir].copy()
                df_para_seleccionar.insert(0, "Seleccionar", st.session_state.default_select_val_tab1) 
                columnas_deshabilitadas = [col for col in df_para_seleccionar.columns if col != 'Seleccionar']

                df_editado = st.data_editor(
                    df_para_seleccionar,
                    use_container_width=True,
                    hide_index=True,
                    column_config={
                        "Seleccionar": st.column_config.CheckboxColumn("Seleccionar", default=st.session_state.default_select_val_tab1),
                        "importe": st.column_config.NumberColumn("Importe", format="$ %d"),
                        "dias_vencido": st.column_config.NumberColumn("Días Vencido", format="%d días"),
                        "fecha_documento": st.column_config.DateColumn("Fecha Emisión", format="YYYY-MM-DD"),
                        "dias_emision": st.column_config.NumberColumn("Días Emisión", format="%d días")
                    },
                    disabled=columnas_deshabilitadas, 
                    key=st.session_state.data_editor_key_tab1 
                )
                
                df_seleccionado = df_editado[df_editado["Seleccionar"] == True].copy()
                st.markdown(f"**Facturas seleccionadas: {len(df_seleccionado)}**")

                # Lógica de Descarga Excel (Tab 1)
                if not df_seleccionado.empty:
                    df_subir_excel = pd.DataFrame()
                    df_subir_excel['TIPO_DOCUMENTO'] = df_seleccionado['nit'].apply(get_tipo_doc_from_nit_col)
                    df_subir_excel['DOCUMENTO'] = df_seleccionado['nit']
                    df_subir_excel['TITULO_VALOR'] = df_seleccionado['factura_norm']
                    df_subir_excel['VALOR'] = pd.to_numeric(df_seleccionado['importe'], errors='coerce').fillna(0).astype(int)
                    df_subir_excel['FECHA'] = pd.to_datetime(df_seleccionado['fecha_vencimiento'], errors='coerce').apply(format_date)
                    df_subir_excel['CODIGO_CONSULTA'] = 986638
                    excel_data_subir = to_excel(df_subir_excel)
                    excel_data_informativo = to_excel_informativo(df_seleccionado)
                else:
                    excel_data_subir = b""
                    excel_data_informativo = b""

                col_btn1, col_btn2 = st.columns(2)
                with col_btn1:
                    st.download_button(
                        label="📤 Descargar Excel para CARGA (Sistema)", 
                        data=excel_data_subir, 
                        file_name="1_facturas_a_subir_CARGA.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
                        disabled=df_seleccionado.empty,
                        use_container_width=True
                    )
                with col_btn2:
                    st.download_button(
                        label="📋 Descargar Reporte INFORMATIVO (Detalle)", 
                        data=excel_data_informativo, 
                        file_name="1_facturas_a_subir_DETALLE.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
                        disabled=df_seleccionado.empty,
                        use_container_width=True
                    )

        with tab2:
            st.subheader("Facturas a Exonerar de Covinoc")
            st.markdown("Facturas en **Covinoc** (que no están 'Efectiva', 'Negada' o 'Exonerada') pero **NO** en la Cartera Ferreinox.")
            
            st.markdown("---")
            st.subheader("Indicadores de Gestión")
            
            kpi_col1, kpi_col2, kpi_col3 = st.columns(3)
            try:
                monto_total_exonerar = pd.to_numeric(df_a_exonerar['saldo'], errors='coerce').sum()
                clientes_unicos_exonerar = df_a_exonerar['cliente'].nunique()
            except Exception:
                monto_total_exonerar = 0
                clientes_unicos_exonerar = 0

            kpi_col1.metric(label="Nº Facturas a Exonerar", value=f"{len(df_a_exonerar)}")
            kpi_col2.metric(label="Monto Total a Exonerar", value=f"${monto_total_exonerar:,.0f}")
            kpi_col3.metric(label="Nº Clientes Afectados", value=f"{clientes_unicos_exonerar}")
            st.markdown("---")

            columnas_mostrar_exonerar = ['cliente', 'documento', 'titulo_valor', 'factura_norm', 'saldo', 'fecha', 'vencimiento', 'estado', 'clave_unica']
            columnas_existentes_exonerar = [col for col in columnas_mostrar_exonerar if col in df_a_exonerar.columns]
            
            st.dataframe(df_a_exonerar[columnas_existentes_exonerar], use_container_width=True, hide_index=True)

            if not df_a_exonerar.empty:
                df_exonerar_excel = pd.DataFrame()
                df_exonerar_excel['TIPO_DOCUMENTO'] = df_a_exonerar['documento'].apply(get_tipo_doc_from_nit_col)
                df_exonerar_excel['DOCUMENTO'] = df_a_exonerar['documento']
                df_exonerar_excel['TITULO_VALOR'] = df_a_exonerar['factura_norm']
                df_exonerar_excel['VALOR'] = pd.to_numeric(df_a_exonerar['saldo'], errors='coerce').fillna(0).astype(int)
                df_exonerar_excel['FECHA'] = pd.to_datetime(df_a_exonerar['vencimiento'], errors='coerce').apply(format_date)
                excel_data_exonerar = to_excel(df_exonerar_excel)
            else:
                excel_data_exonerar = b""

            st.download_button(
                label="📥 Descargar Excel para Exoneración (Formato Covinoc)", 
                data=excel_data_exonerar, 
                file_name="2_exoneraciones_totales.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
                disabled=df_a_exonerar.empty 
            )

        with tab3:
            # ======================================================================================
            # --- MODIFICACIÓN CLAVE: LÓGICA DE ALERTAS 70 DÍAS ---
            # ======================================================================================
            st.subheader("Facturas para Aviso de No Pago y Reclamación")
            st.markdown("Facturas que están **en ambos reportes**, tienen **>= 25 días** vencidas. Se ha agregado el cálculo de fecha límite para reclamación.")
            
            st.markdown("---")
            st.subheader("Indicadores de Gestión")

            columnas_alerta_aviso = [
                'fecha_limite_reclamacion',
                'dias_restantes_reclamo',
                'alerta_estado',
                'estado_kpi_norm'
            ]

            for columna_alerta in columnas_alerta_aviso:
                if columna_alerta not in df_aviso_no_pago.columns:
                    df_aviso_no_pago[columna_alerta] = pd.NA
            
            if not df_aviso_no_pago.empty:
                # --- CÁLCULO FECHA LÍMITE (70 DÍAS) ---
                df_aviso_no_pago['fecha_limite_reclamacion'] = pd.to_datetime(
                    df_aviso_no_pago['fecha_vencimiento_cartera'], errors='coerce'
                ) + timedelta(days=70)
                today_ts = pd.Timestamp.now()
                df_aviso_no_pago['dias_restantes_reclamo'] = (df_aviso_no_pago['fecha_limite_reclamacion'] - today_ts).dt.days

                # Categoría Alerta Visual
                def categorizar_alerta(dias):
                    if pd.isna(dias):
                        return "Sin fecha"
                    if dias < 0:
                        return "🔴 VENCIDO"
                    elif dias <= 15:
                        return "🟠 CRÍTICO"
                    elif dias <= 30:
                        return "🟡 ATENCIÓN"
                    else:
                        return "🟢 A TIEMPO"
                
                df_aviso_no_pago['alerta_estado'] = df_aviso_no_pago['dias_restantes_reclamo'].apply(categorizar_alerta)

                df_aviso_no_pago['estado_kpi_norm'] = df_aviso_no_pago['estado_covinoc'].astype(str).str.upper().str.replace(' ', '')
                # Facturas YA en Aviso
                df_para_reclamar = df_aviso_no_pago[
                    df_aviso_no_pago['estado_kpi_norm'].str.contains("AVISO", na=False) 
                ].copy()
                # Facturas para enviar Aviso
                df_para_enviar_aviso = df_aviso_no_pago[
                    ~df_aviso_no_pago['estado_kpi_norm'].str.contains("AVISO", na=False)
                ].copy()
                
                # --- SUBSET PARA DOCS DE RECLAMACIÓN (>= 70 DÍAS Y ESTADO AVISO) ---
                df_docs_reclamacion = df_aviso_no_pago[
                    (df_aviso_no_pago['dias_vencido_cartera'] >= 70) &
                    (df_aviso_no_pago['estado_covinoc'].astype(str).str.upper().str.contains("AVISO"))
                ].copy()
                
                criticos_reclamacion = len(df_aviso_no_pago[df_aviso_no_pago['dias_restantes_reclamo'] <= 15])

            else:
                df_para_reclamar = pd.DataFrame(columns=df_aviso_no_pago.columns)
                df_para_enviar_aviso = pd.DataFrame(columns=df_aviso_no_pago.columns)
                df_docs_reclamacion = pd.DataFrame(columns=df_aviso_no_pago.columns)
                criticos_reclamacion = 0

            kpi_col1, kpi_col2, kpi_col3 = st.columns(3)
            try:
                total_facturas_aviso = len(df_aviso_no_pago)
                monto_total_aviso = pd.to_numeric(df_aviso_no_pago['importe_cartera'], errors='coerce').sum()
                monto_docs_reclamacion = pd.to_numeric(df_docs_reclamacion['importe_cartera'], errors='coerce').sum()
            except Exception:
                monto_total_aviso = 0
                total_facturas_aviso = 0
                monto_docs_reclamacion = 0

            kpi_col1.metric(label="Nº Facturas Totales en Aviso", value=f"{total_facturas_aviso}")
            kpi_col2.metric(label="Monto Total en Aviso", value=f"${monto_total_aviso:,.0f}")
            kpi_col3.metric(
                label="⚠️ Críticos (<15 días para vencer)", 
                value=f"{criticos_reclamacion}", 
                delta_color="inverse"
            )
            
            st.markdown("---")

            # ======================================================================================
            # --- SECCIÓN: GENERACIÓN DE DOCUMENTOS (WORD) ---
            # ======================================================================================
            st.subheader("📂 Generación de Documentos para Reclamación (Diseño Profesional)")
            st.info("Esta sección genera documentos Word (Endoso, Notificación, Aceptación) listos para imprimir, filtrados por cliente, para facturas con más de 70 días.")

            if df_docs_reclamacion.empty:
                st.warning("No hay facturas que cumplan los criterios para reclamación (>= 70 días y estado Aviso).")
            else:
                col_sel_cli, col_info_cli = st.columns([1, 2])
                
                with col_sel_cli:
                    clientes_reclamacion = sorted(df_docs_reclamacion['nombrecliente_cartera'].dropna().unique())
                    cliente_seleccionado = st.selectbox("Seleccione el Cliente a Reclamar:", options=clientes_reclamacion)
                
                if cliente_seleccionado:
                    # Datos del cliente seleccionado
                    df_cli_sel = df_docs_reclamacion[df_docs_reclamacion['nombrecliente_cartera'] == cliente_seleccionado].copy()
                    
                    # Intentar obtener datos extra del cliente
                    nit_cli = df_cli_sel['nit_cartera'].iloc[0] if not df_cli_sel.empty else "N/A"
                    ciudad_cli = df_cli_sel['poblacion'].iloc[0] if 'poblacion' in df_cli_sel.columns else "Pereira" 
                    dir_cli = "Dirección registrada en RUT" 
                    
                    with col_info_cli:
                        st.write(f"**NIT:** {nit_cli}")
                        st.write(f"**Facturas a procesar:** {len(df_cli_sel)}")
                        st.write(f"**Total a Reclamar:** ${pd.to_numeric(df_cli_sel['importe_cartera'], errors='coerce').sum():,.0f}")

                    # Preparar datos para la función generadora
                    datos_facturas = []
                    for _, row in df_cli_sel.iterrows():
                        datos_facturas.append({
                            'factura': row['factura_norm_cartera'],
                            'valor': pd.to_numeric(row['importe_cartera'], errors='coerce'),
                            'fecha_venc': pd.to_datetime(row['fecha_vencimiento_cartera']).strftime('%Y-%m-%d') if pd.notna(row['fecha_vencimiento_cartera']) else "N/A"
                        })

                    # Botón Generar
                    if st.button(f"📄 Generar Documentos para {cliente_seleccionado}"):
                        zip_bytes = generar_documentos_reclamacion(
                            cliente_nombre=cliente_seleccionado,
                            cliente_nit=nit_cli,
                            cliente_dir=dir_cli,
                            cliente_ciudad=ciudad_cli,
                            facturas_data=datos_facturas
                        )
                        
                        st.success("¡Documentos generados correctamente!")
                        st.download_button(
                            label="📥 Descargar Paquete Documental (ZIP)",
                            data=zip_bytes.getvalue(),
                            file_name=f"Reclamacion_{normalizar_nit_simple(str(nit_cli))}.zip",
                            mime="application/zip"
                        )
            
            st.markdown("---")
            # ======================================================================================

            st.write("Facturas que cumplen los criterios (>= 25 días, > 0, no exoneradas, no negadas). **Ordenadas por urgencia de reclamación.**")
            
            opcion_vista = st.radio(
                "Seleccione la vista:",
                (
                    f"Todas las facturas ({len(df_aviso_no_pago)})", 
                    f"Facturas para ENVIAR Aviso ({len(df_para_enviar_aviso)})", 
                    f"Facturas para RECLAMAR (Ya en Aviso) ({len(df_para_reclamar)})"
                ),
                horizontal=True,
                key="radio_aviso_no_pago"
            )

            if "ENVIAR" in opcion_vista:
                df_aviso_display = df_para_enviar_aviso
            elif "RECLAMAR" in opcion_vista:
                df_aviso_display = df_para_reclamar
            else:
                df_aviso_display = df_aviso_no_pago

            # Ordenar por días restantes para priorizar
            if 'dias_restantes_reclamo' in df_aviso_display.columns:
                df_aviso_display = df_aviso_display.sort_values(
                    by='dias_restantes_reclamo',
                    ascending=True,
                    na_position='last'
                )

            columnas_mostrar_aviso = [
                'alerta_estado', 'dias_restantes_reclamo', 'fecha_limite_reclamacion', # Nuevas columnas
                'nombrecliente_cartera', 'nit_cartera', 'factura_norm_cartera', 'fecha_vencimiento_cartera', 'dias_vencido_cartera', 
                'importe_cartera', 'nomvendedor_cartera', 'saldo_covinoc', 'estado_covinoc'
            ]
            
            columnas_existentes_aviso = [col for col in columnas_mostrar_aviso if col in df_aviso_display.columns]
            
            st.dataframe(
                df_aviso_display[columnas_existentes_aviso], 
                use_container_width=True, 
                hide_index=True,
                column_config={
                    "alerta_estado": st.column_config.TextColumn("Estado Límite", width="medium"),
                    "dias_restantes_reclamo": st.column_config.ProgressColumn(
                        "Días Restantes (70)", 
                        format="%d días", 
                        min_value=-20, 
                        max_value=120,
                        help="Días que faltan para cumplir los 70 días desde vencimiento (Límite legal)."
                    ),
                    "fecha_limite_reclamacion": st.column_config.DateColumn("Fecha Límite", format="YYYY-MM-DD"),
                    "importe_cartera": st.column_config.NumberColumn("Valor", format="$ %d")
                }
            )

            # Lógica de Descarga Excel (Tab 3)
            if not df_para_enviar_aviso.empty:
                df_aviso_excel = pd.DataFrame()
                df_aviso_excel['TIPO_DOCUMENTO'] = df_para_enviar_aviso['documento'].apply(get_tipo_doc_from_nit_col)
                df_aviso_excel['DOCUMENTO'] = df_para_enviar_aviso['documento']
                df_aviso_excel['TITULO_VALOR'] = df_para_enviar_aviso['factura_norm_cartera']
                df_aviso_excel['VALOR'] = pd.to_numeric(df_para_enviar_aviso['importe_cartera'], errors='coerce').fillna(0).astype(int)
                df_aviso_excel['FECHA'] = pd.to_datetime(df_para_enviar_aviso['fecha_vencimiento_cartera'], errors='coerce').apply(format_date)
                excel_data_aviso = to_excel(df_aviso_excel)
            else:
                excel_data_aviso = b""

            st.download_button(
                label="📥 Descargar Excel para Aviso de No Pago (SÓLO PARA ENVIAR)", 
                data=excel_data_aviso, 
                file_name="3_aviso_no_pago_PARA_ENVIAR.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
                disabled=df_para_enviar_aviso.empty 
            )
            
            st.markdown("---")
            st.subheader("🚀 Gestión de Avisos por Vendedor (WhatsApp)")
            
            if df_para_enviar_aviso.empty:
                st.info("No hay facturas 'para enviar' Aviso de No Pago.")
            else:
                st.info("Seleccione los vendedores para preparar los mensajes de gestión (para facturas que AÚN NO están en Aviso).")
                
                vendedores_unicos = sorted(df_para_enviar_aviso['nomvendedor_cartera'].dropna().unique())
                vendedores_seleccionados = st.multiselect(
                    "Vendedores a gestionar:", 
                    options=vendedores_unicos, 
                    default=[]
                )

                if not vendedores_seleccionados:
                    st.write("Seleccione uno o más vendedores para continuar.")
                else:
                    df_aviso_filtrado = df_para_enviar_aviso[
                        df_para_enviar_aviso['nomvendedor_cartera'].isin(vendedores_seleccionados)
                    ].copy()
                    
                    grouped = df_aviso_filtrado.groupby('nomvendedor_cartera')
                    
                    for vendor_name, group_df in grouped:
                        st.markdown(f"---")
                        st.markdown(f"#### Vendedor: **{vendor_name}** ({len(group_df)} facturas)")
                        
                        vendor_name_norm = normalizar_nombre(vendor_name)
                        phone_encontrado = VENDEDORES_WHATSAPP.get(vendor_name_norm, "")
                        
                        col1, col2 = st.columns([0.4, 0.6])
                        
                        with col1:
                            phone_manual = st.text_input(
                                "Teléfono (Ej: +57311...):", 
                                value=phone_encontrado, 
                                key=f"phone_{vendor_name_norm}"
                            )
                        
                            # Construir el mensaje
                            mensaje_header = f"Buen día compañero☀🌈\n\nPor favor gestionar la siguiente cartera que presenta más de 20 días vencidos y se encuentra próxima a:\nAVISO DE NO PAGO EN COVINOC 😨⚠\n"
                            
                            mensaje_clientes_facturas = []
                            grouped_by_client = group_df.groupby('nombrecliente_cartera')
                            
                            for client_name, client_df in grouped_by_client:
                                cliente_str = str(client_name).strip()
                                mensaje_clientes_facturas.append(f"\n* Cliente: {cliente_str}")
                                
                                for _, row in client_df.iterrows():
                                    factura = str(row['factura_norm_cartera']).strip()
                                    try:
                                        valor = float(row['importe_cartera'])
                                        valor_str = f"${valor:,.0f}"
                                    except (ValueError, TypeError):
                                        valor_str = str(row['importe_cartera'])
                                    dias = row['dias_vencido_cartera']
                                    
                                    mensaje_clientes_facturas.append(f"    - Factura: {factura} | Valor: {valor_str} | Días Vencidos: {dias}")

                            mensaje_footer = "\n\nAgradecemos indicar novedad o gestión de pago, en caso contrario se avanzará con el proceso de aviso de no pago.\n\nQuedamos pendientes, muchas gracias"
                            mensaje_completo = mensaje_header + "\n".join(mensaje_clientes_facturas) + mensaje_footer
                            
                            phone_limpio = phone_manual.replace(' ', '').replace('+', '').strip()
                            if phone_limpio and not phone_limpio.startswith("57"):
                                    phone_limpio = f"57{phone_limpio}" 

                            mensaje_url_encoded = urllib.parse.quote_plus(mensaje_completo)
                            url_whatsapp = f"https://wa.me/{phone_limpio}?text={mensaje_url_encoded}"
                            
                            with col2:
                                st.write(" ") 
                                st.link_button(
                                    "📲 Enviar a WhatsApp (Web/App)", 
                                    url_whatsapp, 
                                    use_container_width=True, 
                                    disabled=(not phone_manual)
                                )
                            
                            with st.expander("Ver detalle de facturas y mensaje completo"):
                                st.dataframe(group_df[columnas_existentes_aviso], use_container_width=True, hide_index=True)
                                st.text_area(
                                    "Mensaje a Enviar:", 
                                    value=mensaje_completo, 
                                    height=300, 
                                    key=f"msg_{vendor_name_norm}",
                                    disabled=True
                                )

        with tab4:
            st.subheader("Facturas en Reclamación (Informativo)")
            st.markdown("Facturas que figuran en Covinoc con estado **'Reclamada'**.")

            st.markdown("---")
            st.subheader("Indicadores de Gestión")
            
            kpi_col1, kpi_col2, kpi_col3 = st.columns(3)
            try:
                monto_total_reclamadas = pd.to_numeric(df_reclamadas['saldo'], errors='coerce').sum()
                clientes_unicos_reclamadas = df_reclamadas['cliente'].nunique()
            except Exception:
                monto_total_reclamadas = 0
                clientes_unicos_reclamadas = 0

            kpi_col1.metric(label="Nº Facturas Reclamadas", value=f"{len(df_reclamadas)}")
            kpi_col2.metric(label="Monto Total Reclamado", value=f"${monto_total_reclamadas:,.0f}")
            kpi_col3.metric(label="Nº Clientes", value=f"{clientes_unicos_reclamadas}")
            st.markdown("---")
            
            columnas_mostrar_reclamadas = ['cliente', 'documento', 'titulo_valor', 'factura_norm', 'saldo', 'fecha', 'vencimiento', 'estado', 'clave_unica']
            columnas_existentes_reclamadas = [col for col in columnas_mostrar_reclamadas if col in df_reclamadas.columns]
            
            st.dataframe(df_reclamadas[columnas_existentes_reclamadas], use_container_width=True, hide_index=True)

        with tab5:
            st.subheader("Ajustes por Abonos Parciales")
            st.markdown("Facturas en **ambos reportes** donde el **Saldo Covinoc es MAYOR** al **Importe Cartera** (implica un abono no reportado).")
            
            st.markdown("---")
            st.subheader("Indicadores de Gestión")
            
            kpi_col1, kpi_col2, kpi_col3 = st.columns(3)
            try:
                monto_total_ajuste = pd.to_numeric(df_ajustes['diferencia'], errors='coerce').sum()
                clientes_unicos_ajuste = df_ajustes['nombrecliente_cartera'].nunique()
            except Exception:
                monto_total_ajuste = 0
                clientes_unicos_ajuste = 0

            kpi_col1.metric(label="Nº Facturas para Ajuste", value=f"{len(df_ajustes)}")
            kpi_col2.metric(label="Monto Total a Ajustar", value=f"${monto_total_ajuste:,.0f}")
            kpi_col3.metric(label="Nº Clientes Afectados", value=f"{clientes_unicos_ajuste}")
            st.markdown("---")

            columnas_mostrar_ajustes = [
                'nombrecliente_cartera', 'nit_cartera', 'factura_norm_cartera', 'importe_cartera', 
                'saldo_covinoc', 'diferencia', 'dias_vencido_cartera', 'estado_covinoc', 'clave_unica'
            ]
            columnas_existentes_ajustes = [col for col in columnas_mostrar_ajustes if col in df_ajustes.columns]
            
            df_ajustes_display = df_ajustes[columnas_existentes_ajustes].copy()
            for col_moneda in ['importe_cartera', 'saldo_covinoc', 'diferencia']:
                if col_moneda in df_ajustes_display.columns:
                    df_ajustes_display[col_moneda] = df_ajustes_display[col_moneda].map('${:,.0f}'.format)
            
            st.dataframe(df_ajustes_display, use_container_width=True, hide_index=True)
            
            if not df_ajustes.empty:
                df_ajustes_excel = pd.DataFrame()
                df_ajustes_excel['TIPO_DOCUMENTO'] = df_ajustes['documento'].apply(get_tipo_doc_from_nit_col)
                df_ajustes_excel['DOCUMENTO'] = df_ajustes['documento']
                df_ajustes_excel['TITULO_VALOR'] = df_ajustes['factura_norm_cartera']
                df_ajustes_excel['VALOR'] = pd.to_numeric(df_ajustes['diferencia'], errors='coerce').fillna(0).astype(int)
                df_ajustes_excel['FECHA'] = pd.to_datetime(df_ajustes['fecha_vencimiento_cartera'], errors='coerce').apply(format_date)
                excel_data_ajustes = to_excel(df_ajustes_excel)
            else:
                excel_data_ajustes = b""

            st.download_button(
                label="📥 Descargar Excel de Ajuste (Exoneración Parcial)", 
                data=excel_data_ajustes, 
                file_name="5_ajustes_exoneracion_parcial.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
                disabled=df_ajustes.empty
            )

        with tab6:
            st.subheader("Clientes con FAU Digital Pendiente")
            st.markdown("Cruce entre la **cartera actual** y el archivo **reporteCupos** para identificar, por vendedor, los clientes que aún no tienen **FAU_DIGITAL** diligenciado.")
            st.info("Puede subir el archivo manualmente o dejar que el sistema lo busque automáticamente en el entorno local o en Dropbox.")

            uploaded_reporte_cupos = st.file_uploader(
                "Cargar reporteCupos",
                type=['xlsx', 'xls'],
                key='reporte_cupos_uploader',
                help="Archivo de cupos con las columnas DOCUMENTO, NOMBRES y FAU_DIGITAL, entre otras."
            )

            df_reporte_cupos_raw, fuente_reporte_cupos, error_reporte_cupos = obtener_reporte_cupos_df(uploaded_reporte_cupos)

            if fuente_reporte_cupos:
                st.success(f"Fuente detectada: {fuente_reporte_cupos}")
            elif error_reporte_cupos:
                st.warning(error_reporte_cupos)

            if df_reporte_cupos_raw.empty:
                st.info("Cuando el archivo esté disponible, aquí se mostrará automáticamente la gestión por vendedor con los clientes que deben actualizar su FAU_DIGITAL.")
            else:
                try:
                    df_reporte_cupos = preparar_reporte_cupos(df_reporte_cupos_raw)
                    df_fau_pendiente, df_fau_no_relacionados = construir_reporte_fau_pendiente(df_cartera_full, df_reporte_cupos)
                    df_fau_consolidado = pd.concat([df_fau_pendiente, df_fau_no_relacionados], ignore_index=True, sort=False)
                except ValueError as e:
                    st.error(str(e))
                    df_fau_pendiente = pd.DataFrame()
                    df_fau_no_relacionados = pd.DataFrame()
                    df_fau_consolidado = pd.DataFrame()

                if df_fau_consolidado.empty:
                    st.warning("No se encontraron clientes de la cartera actual con FAU_DIGITAL pendiente en el archivo reporteCupos.")
                    if not df_fau_no_relacionados.empty:
                        st.caption(f"Clientes con FAU pendiente en reporteCupos pero sin cruce con cartera actual: {len(df_fau_no_relacionados)}")
                else:
                    if st.session_state.get('acceso_general', False):
                        vendedores_fau = ['Todos'] + sorted(df_fau_consolidado['Vendedor'].dropna().unique().tolist())
                        vendedor_seleccionado = st.selectbox(
                            'Filtrar por vendedor:',
                            options=vendedores_fau,
                            index=0,
                            key='filtro_vendedor_fau'
                        )
                        if vendedor_seleccionado == 'Todos':
                            df_fau_visible = df_fau_consolidado.copy()
                        else:
                            df_fau_visible = df_fau_consolidado[df_fau_consolidado['Vendedor'] == vendedor_seleccionado].copy()
                    else:
                        vendedor_autenticado_norm = normalizar_nombre(st.session_state.get('vendedor_autenticado', ''))
                        df_fau_visible = df_fau_consolidado[
                            df_fau_consolidado['vendedor_norm'] == vendedor_autenticado_norm
                        ].copy()
                        st.caption(f"Vista filtrada para el vendedor autenticado: {st.session_state.get('vendedor_autenticado', 'S/N')}")

                    if df_fau_visible.empty:
                        st.warning("No hay clientes con FAU_DIGITAL pendiente para el filtro seleccionado.")
                    else:
                        st.markdown("---")
                        st.subheader("Indicadores de Gestión")
                        total_clientes_fau = df_fau_visible['Documento Reporte'].nunique()
                        total_vendedores_fau = df_fau_visible['Vendedor'].nunique()
                        saldo_cartera_fau = pd.to_numeric(df_fau_visible['Saldo Cartera'], errors='coerce').sum()
                        cupo_disponible_fau = pd.to_numeric(df_fau_visible['Cupo Disponible'], errors='coerce').sum()

                        kpi_col1, kpi_col2, kpi_col3, kpi_col4 = st.columns(4)
                        kpi_col1.metric('Clientes Pendientes', f"{total_clientes_fau}")
                        kpi_col2.metric('Vendedores Involucrados', f"{total_vendedores_fau}")
                        kpi_col3.metric('Saldo Cartera Relacionado', f"${saldo_cartera_fau:,.0f}")
                        kpi_col4.metric('Cupo Disponible Reporte', f"${cupo_disponible_fau:,.0f}")

                        st.markdown("---")
                        st.subheader("Resumen por Vendedor")
                        df_resumen_vendedor = df_fau_visible.groupby('Vendedor').agg(
                            Clientes_Pendientes=('Documento Reporte', 'nunique'),
                            Saldo_Cartera=('Saldo Cartera', 'sum'),
                            Cupo_Disponible=('Cupo Disponible', 'sum')
                        ).reset_index().sort_values(by=['Clientes_Pendientes', 'Saldo_Cartera'], ascending=[False, False])

                        st.dataframe(
                            df_resumen_vendedor,
                            use_container_width=True,
                            hide_index=True,
                            column_config={
                                'Saldo_Cartera': st.column_config.NumberColumn('Saldo Cartera', format='$ %d'),
                                'Cupo_Disponible': st.column_config.NumberColumn('Cupo Disponible', format='$ %d')
                            }
                        )

                        st.markdown("---")
                        st.subheader("Detalle de Clientes por Vendedor")
                        columnas_detalle_fau = [
                            'Vendedor', 'Cliente', 'NIT Cartera', 'Documento Reporte', 'Tipo Documento', 'Estado Cupo', 'Tipo Firma',
                            'FAU Digital', 'Pagare Digital', 'Cupo Asignado', 'Extracupo', 'Cupo Disponible',
                            'Saldo Cartera', 'Facturas Activas', 'Max Dias Vencido', 'Sucursal', 'Alerta',
                            'Fecha Ultima Factura', 'Fecha Apertura Cupo', 'Registros Reporte', 'Relacion Vendedor'
                        ]
                        columnas_visibles_fau = [col for col in columnas_detalle_fau if col in df_fau_visible.columns]

                        df_fau_export = df_fau_visible[columnas_visibles_fau].copy()
                        st.dataframe(
                            df_fau_export,
                            use_container_width=True,
                            hide_index=True,
                            column_config={
                                'Cupo Asignado': st.column_config.NumberColumn('Cupo Asignado', format='$ %d'),
                                'Extracupo': st.column_config.NumberColumn('Extracupo', format='$ %d'),
                                'Cupo Disponible': st.column_config.NumberColumn('Cupo Disponible', format='$ %d'),
                                'Saldo Cartera': st.column_config.NumberColumn('Saldo Cartera', format='$ %d'),
                                'Fecha Ultima Factura': st.column_config.DateColumn('Fecha Ultima Factura', format='YYYY-MM-DD'),
                                'Fecha Apertura Cupo': st.column_config.DateColumn('Fecha Apertura Cupo', format='YYYY-MM-DD')
                            }
                        )

                        # Exportar el Excel con todos los tipos de documento
                        excel_fau_pendiente = to_excel_fau_pendiente(df_fau_export.drop(columns=['vendedor_norm'], errors='ignore'))
                        st.download_button(
                            label='📥 Descargar Reporte FAU Digital Pendiente',
                            data=excel_fau_pendiente,
                            file_name='6_fau_digital_pendiente.xlsx',
                            mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                            use_container_width=True
                        )

                    if not df_fau_no_relacionados.empty:
                        with st.expander(f"Ver clientes con FAU pendiente sin cruce exacto de cartera ({len(df_fau_no_relacionados)})"):
                            st.dataframe(df_fau_no_relacionados, use_container_width=True, hide_index=True)

        # ======================================================================
        # --- TAB 7: DASHBOARD ESTRATÉGICO DE KPIs (reporteTransacciones) ---
        # ======================================================================
        with tab7:
            try:
                st.subheader("📊 Dashboard Estratégico Covinoc")
                st.markdown("Análisis integral del archivo **reporteTransacciones**: bolsa de garantía, evolución mensual/anual y radiografía de clientes.")

                if dfa_covinoc.empty:
                    st.warning("No hay datos de reporteTransacciones para analizar.")
                else:
                    # ---------- Cálculos base ----------
                    df_res_mes = resumen_mensual_covinoc(dfa_covinoc)
                    acumulado_bolsa = float(dfa_covinoc.loc[dfa_covinoc['fecha_dt'] >= FECHA_INICIO_BOLSA, 'valor_garantizado_num'].sum())
                    pct_bolsa = (acumulado_bolsa / BOLSA_COVINOC * 100) if BOLSA_COVINOC else 0
                    disponible_bolsa = BOLSA_COVINOC - acumulado_bolsa
                    num_meses = max(len(df_res_mes), 1)
                    ritmo_mensual = acumulado_bolsa / num_meses

                    # Selector de mes a analizar
                    meses_disponibles = df_res_mes['periodo'].tolist() if not df_res_mes.empty else []
                    col_sel, col_info = st.columns([1, 2])
                    with col_sel:
                        if meses_disponibles:
                            mes_sel = st.selectbox(
                                "📅 Mes a analizar:",
                                options=meses_disponibles[::-1],
                                format_func=_etiqueta_periodo,
                                key="covinoc_mes_sel"
                            )
                        else:
                            mes_sel = None
                    with col_info:
                        st.info(f"Bolsa de garantía: **${BOLSA_COVINOC:,.0f}** · Conteo desde **{FECHA_INICIO_BOLSA.strftime('%d-%b-%Y')}** · Datos hasta hoy.")

                    # Valor del mes seleccionado y delta vs mes anterior
                    valor_mes = 0.0
                    titulos_mes = 0
                    clientes_mes = 0
                    delta_mes = None
                    if mes_sel and not df_res_mes.empty:
                        fila_mes = df_res_mes[df_res_mes['periodo'] == mes_sel]
                        if not fila_mes.empty:
                            valor_mes = float(fila_mes['valor'].iloc[0])
                            titulos_mes = int(fila_mes['titulos'].iloc[0])
                            clientes_mes = int(fila_mes['clientes'].iloc[0])
                            idx_mes = df_res_mes.index.get_loc(fila_mes.index[0])
                            if idx_mes > 0:
                                valor_prev = float(df_res_mes.iloc[idx_mes - 1]['valor'])
                                delta_mes = valor_mes - valor_prev

                    st.markdown("### 🎯 Indicadores de la Bolsa de Garantía")
                    k1, k2, k3, k4 = st.columns(4)
                    k1.metric(
                        f"Reportado en {_etiqueta_periodo(mes_sel) if mes_sel else 'el mes'}",
                        f"${valor_mes:,.0f}",
                        delta=f"${delta_mes:,.0f} vs mes anterior" if delta_mes is not None else None
                    )
                    k2.metric("Acumulado desde 30-Sep-2025", f"${acumulado_bolsa:,.0f}", delta=f"{pct_bolsa:.1f}% de la bolsa")
                    k3.metric(
                        "Bolsa disponible" if disponible_bolsa >= 0 else "Bolsa sobregirada",
                        f"${disponible_bolsa:,.0f}",
                        delta="Dentro del límite" if disponible_bolsa >= 0 else "Supera el límite",
                        delta_color="normal" if disponible_bolsa >= 0 else "inverse"
                    )
                    k4.metric("Ritmo mensual promedio", f"${ritmo_mensual:,.0f}", delta=f"{num_meses} meses de operación")

                    # ---------- Medidor único de ocupación de la bolsa ----------
                    mg1, mg2, mg3 = st.columns([1, 2, 1])
                    with mg2:
                        fig_gauge = go.Figure(go.Indicator(
                            mode="gauge+number",
                            value=acumulado_bolsa,
                            number={'prefix': "$", 'valueformat': ',.0f', 'font': {'size': 32}},
                            title={'text': "<b>Ocupación acumulada de la bolsa</b><br><span style='font-size:12px'>desde 30-Sep-2025 · no se libera con pagos</span>"},
                            gauge={
                                'axis': {'range': [0, max(BOLSA_COVINOC, acumulado_bolsa)], 'tickformat': ',.0f'},
                                'bar': {'color': PALETA_COLORES['primario']},
                                'steps': [
                                    {'range': [0, BOLSA_COVINOC * 0.7], 'color': '#E8F5E9'},
                                    {'range': [BOLSA_COVINOC * 0.7, BOLSA_COVINOC * 0.9], 'color': '#FEF4C0'},
                                    {'range': [BOLSA_COVINOC * 0.9, max(BOLSA_COVINOC, acumulado_bolsa)], 'color': '#FDE0DE'},
                                ],
                                'threshold': {'line': {'color': PALETA_COLORES['secundario'], 'width': 4}, 'thickness': 0.85, 'value': BOLSA_COVINOC}
                            }
                        ))
                        fig_gauge.update_layout(height=320, margin=dict(l=20, r=20, t=70, b=10), paper_bgcolor='rgba(0,0,0,0)')
                        st.plotly_chart(fig_gauge, use_container_width=True)
                        st.caption(f"🔴 Línea roja = límite de la bolsa (${BOLSA_COVINOC:,.0f}). Ocupación actual: **{pct_bolsa:.1f}%**.")

                    # ---------- Evolución mensual ----------
                    st.markdown("---")
                    st.markdown("### 📈 Evolución mensual del valor garantizado")
                    if not df_res_mes.empty:
                        fig_mes = go.Figure()
                        fig_mes.add_trace(go.Bar(
                            x=df_res_mes['mes_label'], y=df_res_mes['valor'],
                            name="Valor del mes", marker_color=PALETA_COLORES['primario'],
                            text=[f"${v/1e6:,.0f}M" for v in df_res_mes['valor']], textposition='outside'
                        ))
                        fig_mes.add_trace(go.Scatter(
                            x=df_res_mes['mes_label'], y=df_res_mes['acumulado'],
                            name="Acumulado", mode='lines+markers', yaxis='y2',
                            line=dict(color=PALETA_COLORES['acento'], width=3)
                        ))
                        fig_mes.add_hline(y=BOLSA_COVINOC, line_dash="dash", line_color=PALETA_COLORES['secundario'],
                                          annotation_text=f"Bolsa ${BOLSA_COVINOC/1e6:,.0f}M", annotation_position="top left", yref='y2')
                        fig_mes.update_layout(
                            height=420, hovermode='x unified', plot_bgcolor='rgba(0,0,0,0)', paper_bgcolor='rgba(0,0,0,0)',
                            legend=dict(orientation='h', yanchor='bottom', y=1.02, xanchor='right', x=1),
                            yaxis=dict(title="Valor del mes", tickformat=',.0f'),
                            yaxis2=dict(title="Acumulado", overlaying='y', side='right', tickformat=',.0f', showgrid=False),
                            margin=dict(l=20, r=20, t=40, b=20)
                        )
                        st.plotly_chart(fig_mes, use_container_width=True)

                        df_tabla_mes = df_res_mes[['mes_label', 'valor', 'titulos', 'clientes', 'acumulado']].copy()
                        df_tabla_mes.columns = ['Mes', 'Valor Garantizado', 'Títulos', 'Clientes', 'Acumulado Bolsa']
                        with st.expander("Ver tabla mensual detallada"):
                            st.dataframe(
                                df_tabla_mes, use_container_width=True, hide_index=True,
                                column_config={
                                    'Valor Garantizado': st.column_config.NumberColumn(format='$ %d'),
                                    'Acumulado Bolsa': st.column_config.NumberColumn(format='$ %d'),
                                }
                            )

                    # ---------- Radiografía de clientes ----------
                    st.markdown("---")
                    st.markdown("### 👥 Radiografía de Clientes")

                    clientes_con_facturas = int(dfa_covinoc['documento'].nunique())
                    clientes_con_aviso = int(dfa_covinoc.loc[dfa_covinoc['tiene_aviso'], 'documento'].nunique())
                    # Activos vencidos
                    hoy_dt = pd.to_datetime(datetime.now().date())
                    mask_venc = dfa_covinoc['es_activo'] & dfa_covinoc['vencimiento_dt'].notna() & (dfa_covinoc['vencimiento_dt'] < hoy_dt)
                    clientes_vencidos = int(dfa_covinoc.loc[mask_venc, 'documento'].nunique())
                    titulos_vencidos = int(mask_venc.sum())
                    saldo_vencido = float(dfa_covinoc.loc[mask_venc, 'saldo_num'].sum())

                    # Cupos
                    df_cupo_uso = analizar_cupos_clientes(df_cupos_auto, dfa_covinoc)
                    if not df_cupo_uso.empty:
                        creados_cupo = int(df_cupo_uso['documento_norm'].nunique())
                        con_cupo_asignado = int(df_cupo_uso.loc[df_cupo_uso['tiene_cupo'], 'documento_norm'].nunique())
                        usan_cupo = int(df_cupo_uso.loc[df_cupo_uso['usa_cupo'], 'documento_norm'].nunique())
                        no_usan_cupo = int(df_cupo_uso.loc[df_cupo_uso['tiene_cupo'] & (~df_cupo_uso['usa_cupo']), 'documento_norm'].nunique())
                        cupo_asignado_total = float(df_cupo_uso['cupo_asignado'].sum())
                        cupo_disponible_total = float(df_cupo_uso['cupo_disponible'].sum())
                    else:
                        creados_cupo = con_cupo_asignado = usan_cupo = no_usan_cupo = 0
                        cupo_asignado_total = cupo_disponible_total = 0.0

                    r1, r2, r3, r4 = st.columns(4)
                    r1.metric("Clientes con facturas cargadas", f"{clientes_con_facturas}")
                    r2.metric("Clientes creados con cupo", f"{creados_cupo}" if creados_cupo else "—",
                              help="Requiere el archivo reporteCupos" if not creados_cupo else None)
                    r3.metric("Clientes que USAN el cupo", f"{usan_cupo}" if creados_cupo else "—")
                    r4.metric("Clientes que NO usan el cupo", f"{no_usan_cupo}" if creados_cupo else "—",
                              delta="Oportunidad de activación" if no_usan_cupo else None, delta_color="inverse")

                    r5, r6, r7, r8 = st.columns(4)
                    r5.metric("Clientes con avisos de no pago", f"{clientes_con_aviso}")
                    r6.metric("Clientes con títulos vencidos", f"{clientes_vencidos}")
                    r7.metric("Títulos activos vencidos", f"{titulos_vencidos}", delta=f"saldo ${saldo_vencido:,.0f}", delta_color="inverse")
                    r8.metric("Cupo total disponible", f"${cupo_disponible_total:,.0f}" if creados_cupo else "—",
                              delta=f"de ${cupo_asignado_total:,.0f} asignado" if creados_cupo else None)

                    if not creados_cupo:
                        st.caption("💡 Carga el archivo **reporteCupos** (en la pestaña 6 o en la pestaña 8) para completar los indicadores de cupo.")

                    # ---------- Distribuciones y rankings ----------
                    st.markdown("---")
                    st.markdown("### 🔎 Análisis y Rankings")
                    d1, d2 = st.columns(2)
                    with d1:
                        st.markdown("**Distribución por estado de los títulos**")
                        df_estado = dfa_covinoc.groupby('estado_norm').agg(
                            titulos=('valor_garantizado_num', 'size'),
                            valor=('valor_garantizado_num', 'sum')
                        ).reset_index().sort_values('valor', ascending=False)
                        fig_estado = px.pie(
                            df_estado, names='estado_norm', values='valor', hole=0.55,
                            color_discrete_sequence=[PALETA_COLORES['primario'], PALETA_COLORES['secundario'],
                                                     PALETA_COLORES['acento'], PALETA_COLORES['destacado'],
                                                     PALETA_COLORES['exito_verde'], '#9E9E9E', '#607D8B']
                        )
                        fig_estado.update_traces(textposition='inside', textinfo='percent+label')
                        fig_estado.update_layout(height=340, showlegend=False, margin=dict(l=10, r=10, t=10, b=10), paper_bgcolor='rgba(0,0,0,0)')
                        st.plotly_chart(fig_estado, use_container_width=True)
                    with d2:
                        st.markdown("**Top 10 clientes por valor garantizado**")
                        df_top_cli = dfa_covinoc.groupby('cliente').agg(
                            valor=('valor_garantizado_num', 'sum'), titulos=('valor_garantizado_num', 'size')
                        ).reset_index().sort_values('valor', ascending=False).head(10)
                        fig_top = px.bar(
                            df_top_cli.sort_values('valor'), x='valor', y='cliente', orientation='h',
                            text=[f"${v/1e6:,.1f}M" for v in df_top_cli.sort_values('valor')['valor']],
                            color_discrete_sequence=[PALETA_COLORES['primario']]
                        )
                        fig_top.update_layout(height=340, plot_bgcolor='rgba(0,0,0,0)', paper_bgcolor='rgba(0,0,0,0)',
                                              xaxis_title=None, yaxis_title=None, margin=dict(l=10, r=10, t=10, b=10))
                        st.plotly_chart(fig_top, use_container_width=True)

                    # ---------- Segmentación de clientes por riesgo ----------
                    st.markdown("---")
                    st.markdown("### 🚦 Segmentación de Clientes por Riesgo")
                    st.caption("Los clientes **con avisos de no pago son de RIESGO** (conviene mantenerlos protegidos con Covinoc). Los que **nunca han tenido avisos son de bajo riesgo**: podrían retirarse de Covinoc para no ocupar bolsa sin necesidad.")

                    clientes_riesgo_set = set(dfa_covinoc.loc[dfa_covinoc['tiene_aviso'], 'cliente'].dropna())

                    df_riesgo = dfa_covinoc[dfa_covinoc['cliente'].isin(clientes_riesgo_set)].groupby('cliente').agg(
                        avisos=('tiene_aviso', 'sum'),
                        titulos=('valor_garantizado_num', 'size'),
                        valor=('valor_garantizado_num', 'sum'),
                    ).reset_index()
                    df_riesgo['avisos'] = df_riesgo['avisos'].astype(int)
                    df_riesgo = df_riesgo.sort_values('avisos', ascending=False)

                    df_buenos = dfa_covinoc[~dfa_covinoc['cliente'].isin(clientes_riesgo_set)].groupby('cliente').agg(
                        titulos=('valor_garantizado_num', 'size'),
                        valor=('valor_garantizado_num', 'sum'),
                        ultima_txn=('fecha_dt', 'max'),
                    ).reset_index().sort_values('valor', ascending=False)

                    seg1, seg2, seg3 = st.columns(3)
                    seg1.metric("🔴 Clientes de RIESGO (con avisos)", f"{df_riesgo['cliente'].nunique()}")
                    seg2.metric("🟢 Clientes BUENOS (0 avisos)", f"{df_buenos['cliente'].nunique()}")
                    seg3.metric("Bolsa ocupada por clientes buenos", f"${df_buenos['valor'].sum():,.0f}",
                                delta="Potencial a liberar", delta_color="inverse")

                    col_r, col_b = st.columns(2)
                    with col_r:
                        st.markdown("**🔴 Clientes de RIESGO — mantener protegidos** (más avisos = más se dejan vencer)")
                        if not df_riesgo.empty:
                            df_r_show = df_riesgo.head(15).rename(columns={
                                'cliente': 'Cliente', 'avisos': 'N° Avisos', 'titulos': 'Títulos', 'valor': 'Valor Garantizado'})
                            st.dataframe(
                                df_r_show, use_container_width=True, hide_index=True,
                                column_config={'Valor Garantizado': st.column_config.NumberColumn(format='$ %d'),
                                               'N° Avisos': st.column_config.ProgressColumn(
                                                   'N° Avisos', format='%d', min_value=0, max_value=int(df_riesgo['avisos'].max()))}
                            )
                        else:
                            st.info("No hay clientes con avisos de no pago.")
                    with col_b:
                        st.markdown("**🟢 Clientes BUENOS — candidatos a retirar de Covinoc** (nunca han tenido avisos)")
                        if not df_buenos.empty:
                            df_b_show = df_buenos.head(15).rename(columns={
                                'cliente': 'Cliente', 'titulos': 'Títulos', 'valor': 'Valor Garantizado', 'ultima_txn': 'Última Transacción'})
                            st.dataframe(
                                df_b_show, use_container_width=True, hide_index=True,
                                column_config={'Valor Garantizado': st.column_config.NumberColumn(format='$ %d'),
                                               'Última Transacción': st.column_config.DateColumn(format='YYYY-MM-DD')}
                            )
                        else:
                            st.info("Sin datos.")

                    # ---------- Descarga análisis completo ----------
                    st.markdown("---")
                    excel_analisis = to_excel_generico({
                        'Resumen Mensual': df_tabla_mes if not df_res_mes.empty else pd.DataFrame(),
                        'Por Estado': df_estado.rename(columns={'estado_norm': 'Estado', 'titulos': 'Títulos', 'valor': 'Valor'}),
                        'Top Clientes': df_top_cli.rename(columns={'cliente': 'Cliente', 'valor': 'Valor', 'titulos': 'Títulos'}),
                        'Clientes Riesgo': df_riesgo.rename(columns={'cliente': 'Cliente', 'avisos': 'N Avisos', 'titulos': 'Titulos', 'valor': 'Valor'}),
                        'Clientes Buenos (Retirar)': df_buenos.rename(columns={'cliente': 'Cliente', 'titulos': 'Titulos', 'valor': 'Valor', 'ultima_txn': 'Ultima Transaccion'}),
                    })
                    st.download_button(
                        "📥 Descargar análisis completo (Excel)",
                        data=excel_analisis,
                        file_name=f"analisis_covinoc_{datetime.now().strftime('%Y%m%d')}.xlsx",
                        mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                        use_container_width=True
                    )

            except Exception as _err_seccion:
                st.error("⚠️ No se pudo construir esta sección. Las demás pestañas siguen funcionando. Detalle del error:")
                st.exception(_err_seccion)

        # ======================================================================
        # --- TAB 8: ACTIVACIÓN DE CLIENTES (WhatsApp + Correo SendGrid) ---
        # ======================================================================
        with tab8:
            try:
                st.subheader("🚀 Activación de Clientes con Cupo sin Usar")
                st.markdown("Identifica clientes con **cupo aprobado pero sin uso** y lanza campañas de activación por **WhatsApp** (link directo) y **correo masivo** (SendGrid).")

                uploaded_cupos_tab8 = st.file_uploader(
                    "Cargar reporteCupos (opcional, si no se detecta automáticamente)",
                    type=['xlsx', 'xls'], key='reporte_cupos_uploader_tab8'
                )
                if uploaded_cupos_tab8 is not None:
                    try:
                        df_cupos_tab8 = preparar_reporte_cupos(leer_reporte_cupos_excel(uploaded_cupos_tab8))
                        st.success("Archivo reporteCupos cargado manualmente.")
                    except Exception as e:
                        st.error(f"No fue posible leer el archivo: {e}")
                        df_cupos_tab8 = df_cupos_auto
                else:
                    df_cupos_tab8 = df_cupos_auto
                    if fuente_cupos_auto:
                        st.caption(f"Fuente detectada automáticamente: {fuente_cupos_auto}")

                df_cupo_uso8 = analizar_cupos_clientes(df_cupos_tab8, dfa_covinoc)

                if df_cupo_uso8.empty:
                    st.warning("No hay datos de reporteCupos disponibles. Cárgalo arriba para construir la campaña de activación.")
                else:
                    # Clientes con cupo asignado que NO lo usan
                    df_no_usan = df_cupo_uso8[df_cupo_uso8['tiene_cupo'] & (~df_cupo_uso8['usa_cupo'])].copy()

                    # Enriquecer con contacto de la cartera
                    contacto = resumen_contacto_cartera(df_cartera_full)
                    if not contacto.empty:
                        set_nits = set(contacto['nit_norm_cartera'])
                        def _match_nit(doc):
                            if not isinstance(doc, str) or not doc:
                                return None
                            if doc in set_nits:
                                return doc
                            if doc[:-1] in set_nits:
                                return doc[:-1]
                            return None
                        df_no_usan['nit_match'] = df_no_usan['documento_norm'].apply(_match_nit)
                        df_camp = df_no_usan.merge(contacto, left_on='nit_match', right_on='nit_norm_cartera', how='left')
                    else:
                        df_camp = df_no_usan.copy()
                        for c in ['cliente_cartera', 'telefono', 'email', 'vendedor_cartera', 'cod_cliente']:
                            df_camp[c] = ''

                    # Nombre y vendedor consolidados
                    df_camp['cliente_final'] = df_camp['cliente_cartera'].fillna('').replace('', pd.NA)
                    df_camp['cliente_final'] = df_camp['cliente_final'].fillna(df_camp['nombres'])
                    df_camp['vendedor_final'] = df_camp['vendedor_cartera'].fillna('').replace('', pd.NA)
                    df_camp['vendedor_final'] = df_camp['vendedor_final'].fillna('GESTION INTERNA')
                    df_camp['email'] = df_camp['email'].fillna('').astype(str).str.strip()
                    df_camp['telefono'] = df_camp['telefono'].fillna('').astype(str).str.strip()
                    df_camp['email_valido'] = df_camp['email'].str.contains('@', na=False)
                    df_camp['wa_link'] = df_camp.apply(
                        lambda r: generar_link_wa_activacion(r['telefono'], r['cliente_final'], r['cupo_disponible']), axis=1
                    )

                    # KPIs de la campaña
                    total_no_usan = int(df_camp['documento_norm'].nunique())
                    con_wa = int(df_camp['wa_link'].notna().sum())
                    con_email = int(df_camp['email_valido'].sum())
                    cupo_ocioso = float(df_camp['cupo_disponible'].sum())

                    c1, c2, c3, c4 = st.columns(4)
                    c1.metric("Clientes con cupo sin usar", f"{total_no_usan}")
                    c2.metric("Con celular (WhatsApp)", f"{con_wa}")
                    c3.metric("Con correo válido", f"{con_email}")
                    c4.metric("Cupo ocioso (sin activar)", f"${cupo_ocioso:,.0f}")

                    # Filtro por vendedor
                    st.markdown("---")
                    vendedores_camp = ['Todos'] + sorted([v for v in df_camp['vendedor_final'].dropna().unique() if str(v).strip()])
                    vend_sel = st.selectbox("Filtrar por vendedor:", options=vendedores_camp, key="camp_vendedor")
                    df_vista = df_camp if vend_sel == 'Todos' else df_camp[df_camp['vendedor_final'] == vend_sel]

                    canal = st.radio(
                        "Canal de activación:",
                        ["💬 WhatsApp (links directos)", "📧 Correo masivo (SendGrid)"],
                        horizontal=True, key="camp_canal"
                    )

                    # ----- WhatsApp -----
                    if canal.startswith("💬"):
                        st.markdown("Cada cliente tiene un **link de WhatsApp con el mensaje de activación ya redactado**. Haz clic en 'Abrir' para enviarlo desde tu WhatsApp.")
                        df_wa = df_vista[df_vista['wa_link'].notna()][
                            ['cliente_final', 'documento', 'telefono', 'cupo_disponible', 'vendedor_final', 'wa_link']
                        ].copy()
                        df_wa.columns = ['Cliente', 'Documento', 'Teléfono', 'Cupo Disponible', 'Vendedor', 'WhatsApp']
                        if df_wa.empty:
                            st.info("No hay clientes con celular válido para este filtro.")
                        else:
                            st.dataframe(
                                df_wa, use_container_width=True, hide_index=True,
                                column_config={
                                    'Cupo Disponible': st.column_config.NumberColumn(format='$ %d'),
                                    'WhatsApp': st.column_config.LinkColumn('WhatsApp', display_text="💬 Abrir chat")
                                }
                            )
                            st.download_button(
                                "📥 Descargar lista WhatsApp (Excel)",
                                data=to_excel_generico({'WhatsApp Activacion': df_wa}),
                                file_name=f"activacion_whatsapp_{datetime.now().strftime('%Y%m%d')}.xlsx",
                                mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                                use_container_width=True
                            )

                    # ----- Correo masivo SendGrid -----
                    else:
                        df_mail = df_vista[df_vista['email_valido']][
                            ['cliente_final', 'documento', 'email', 'cupo_disponible', 'vendedor_final']
                        ].copy().drop_duplicates(subset=['email'])
                        st.markdown(f"Se enviará **un correo de activación** con diseño institucional a **{len(df_mail)} clientes** con correo válido.")

                        if st.checkbox("👁️ Previsualizar el correo (ejemplo)", key="camp_preview"):
                            ejemplo = df_mail.iloc[0] if not df_mail.empty else None
                            html_prev = plantilla_activacion_html(
                                ejemplo['cliente_final'] if ejemplo is not None else "Cliente Ejemplo",
                                ejemplo['cupo_disponible'] if ejemplo is not None else 1500000,
                                ejemplo['vendedor_final'] if ejemplo is not None else "Ferreinox"
                            )
                            components.html(html_prev, height=620, scrolling=True)

                        asunto_mail = st.text_input(
                            "Asunto del correo:",
                            value="🎉 Tu cupo de crédito Ferreinox está listo para usar",
                            key="camp_asunto"
                        )

                        if "sendgrid" not in st.secrets:
                            st.warning("⚠️ No está configurada la sección **[sendgrid]** en secrets (api_key, from_email, from_name). Sin esto no se puede enviar correo.")
                        else:
                            api_key = st.secrets["sendgrid"].get("api_key", "")
                            from_email = st.secrets["sendgrid"].get("from_email", "")
                            from_name = st.secrets["sendgrid"].get("from_name", "Ferreinox S.A.S. BIC")

                            # ---- ENVÍO DE PRUEBA (siempre disponible) ----
                            st.markdown("---")
                            st.markdown("#### ✉️ 1) Envíate un correo de PRUEBA")
                            col_tp1, col_tp2 = st.columns([3, 1])
                            with col_tp1:
                                correo_prueba = st.text_input(
                                    "Escribe TU correo para recibir la prueba:",
                                    key="camp_prueba", placeholder="tucorreo@ferreinox.co"
                                )
                            with col_tp2:
                                st.markdown("<div style='height:28px'></div>", unsafe_allow_html=True)
                                enviar_prueba = st.button("✉️ Enviar prueba", use_container_width=True)
                            if enviar_prueba:
                                if correo_prueba and '@' in correo_prueba:
                                    if not df_mail.empty:
                                        ej = df_mail.iloc[0]
                                        ej_cli, ej_cupo, ej_vend = ej['cliente_final'], ej['cupo_disponible'], ej['vendedor_final']
                                    else:
                                        ej_cli, ej_cupo, ej_vend = "Cliente Ejemplo", 1500000, "Cartera Ferreinox"
                                    html_test = plantilla_activacion_html(ej_cli, ej_cupo, ej_vend)
                                    with st.spinner(f"Enviando prueba a {correo_prueba}..."):
                                        ok, det = enviar_correo_activacion_sendgrid(
                                            api_key, from_email, from_name, correo_prueba.strip(),
                                            ej_cli, asunto_mail, html_test,
                                            "Tu cupo de crédito Ferreinox está listo para usar."
                                        )
                                    if ok:
                                        st.success(f"✅ Prueba enviada a {correo_prueba}. Revisa tu bandeja (y la carpeta de spam).")
                                    else:
                                        st.error(f"❌ No se pudo enviar: {det}")
                                else:
                                    st.error("Ingresa un correo válido (que tenga @).")

                            # ---- ENVÍO MASIVO ----
                            st.markdown("---")
                            st.markdown("#### 🚀 2) Envío masivo a los clientes")
                            if df_mail.empty:
                                st.info("No hay clientes con correo válido en este filtro para el envío masivo.")
                            else:
                                confirmar = st.checkbox(f"Confirmo el envío masivo a {len(df_mail)} clientes", key="camp_confirmar")
                                if st.button("🚀 Enviar campaña masiva", type="primary", use_container_width=True, disabled=not confirmar):
                                    barra = st.progress(0.0)
                                    estado_envio = st.empty()
                                    enviados, fallidos = 0, 0
                                    errores = []
                                    total = len(df_mail)
                                    for i, (_, fila) in enumerate(df_mail.iterrows()):
                                        html_c = plantilla_activacion_html(fila['cliente_final'], fila['cupo_disponible'], fila['vendedor_final'])
                                        ok, det = enviar_correo_activacion_sendgrid(
                                            api_key, from_email, from_name, fila['email'],
                                            fila['cliente_final'], asunto_mail,
                                            html_c, "Tu cupo de crédito Ferreinox está listo para usar."
                                        )
                                        if ok:
                                            enviados += 1
                                        else:
                                            fallidos += 1
                                            errores.append({'Cliente': fila['cliente_final'], 'Correo': fila['email'], 'Error': det})
                                        barra.progress((i + 1) / total)
                                        estado_envio.caption(f"Enviando... {i + 1}/{total} · ✅ {enviados} · ❌ {fallidos}")
                                    if enviados:
                                        st.success(f"✅ Campaña finalizada: {enviados} correos enviados.")
                                    if fallidos:
                                        st.error(f"❌ {fallidos} correos fallaron.")
                                        st.dataframe(pd.DataFrame(errores), use_container_width=True, hide_index=True)

                        if not df_mail.empty:
                            df_mail_dl = df_mail.rename(columns={
                                'cliente_final': 'Cliente', 'documento': 'Documento', 'email': 'Correo',
                                'cupo_disponible': 'Cupo Disponible', 'vendedor_final': 'Vendedor'
                            })
                            st.download_button(
                                "📥 Descargar lista de correos (Excel)",
                                data=to_excel_generico({'Correos Activacion': df_mail_dl}),
                                file_name=f"activacion_correos_{datetime.now().strftime('%Y%m%d')}.xlsx",
                                mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                                use_container_width=True
                            )


            except Exception as _err_seccion:
                st.error("⚠️ No se pudo construir esta sección. Las demás pestañas siguen funcionando. Detalle del error:")
                st.exception(_err_seccion)

if __name__ == '__main__':
    main()
